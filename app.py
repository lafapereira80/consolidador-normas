import streamlit as st
import tempfile
import io
import json
import os
import re
import time
import copy
import hashlib
import base64
from html.parser import HTMLParser
from html import unescape
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from streamlit.runtime.scriptrunner import add_script_run_ctx, get_script_run_ctx
from google import genai
from google.genai import types
from pydantic import BaseModel, Field, ValidationError
from typing import List, Optional, Any, Dict

from supabase import create_client, Client
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from auth_utils import gerar_hash_senha, verificar_senha
from streamlit_quill import st_quill

try:
    from groq import Groq
except ImportError:
    Groq = None
try:
    from mistralai import Mistral
except ImportError:
    try:
        from mistralai.client import Mistral
    except ImportError:
        Mistral = None
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

# ----------------- HUB MULTI-IA -----------------
PROVEDORES_IA = {
    "Google Gemini": {
        "motor": "gemini",
        "modelos": ["gemini-3.7-flash", "gemini-3.6-flash", "gemini-3.5-flash"],
        "secret": "GEMINI_API_KEY",
    },
    "Groq (Llama / GPT-OSS)": {
        "motor": "groq",
        "modelos": ["openai/gpt-oss-120b", "openai/gpt-oss-20b", "llama-3.1-70b-versatile"],
        "secret": "GROQ_API_KEY",
    },
    "OpenRouter (Qwen / DeepSeek / Llama)": {
        "motor": "openrouter",
        "modelos": ["qwen/qwen-2.5-72b-instruct", "meta-llama/llama-3.3-70b-instruct:free"],
        "secret": "OPENROUTER_API_KEY",
    },
    "Mistral AI (Small / Nemo)": {
        "motor": "mistral",
        "modelos": ["mistral-small-latest", "open-mistral-nemo"],
        "secret": "MISTRAL_API_KEY",
    },
}

def submit_com_contexto(executor, fn, *args, **kwargs):
    ctx = get_script_run_ctx()
    def _wrapper(*a, **kw):
        if ctx is not None:
            add_script_run_ctx(threading.current_thread(), ctx)
        return fn(*a, **kw)
    return executor.submit(_wrapper, *args, **kwargs)

def obter_chave_provedor(nome_provedor):
    cfg = PROVEDORES_IA[nome_provedor]
    try:
        return st.secrets[cfg["secret"]]
    except Exception:
        try:
            return st.secrets["api_keys"][cfg["secret"]]
        except Exception:
            return os.environ.get(cfg["secret"], "")

try:
    from weasyprint import HTML as WeasyHTML, CSS as WeasyCSS
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False

st.set_page_config(page_title="Autopilot Normativo", page_icon="⚖️", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="collapsedControl"] { display: none !important; }
    
    .block-container { padding-top: 2rem; }
    .main-header {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        padding: 30px 20px;
        border-radius: 12px;
        color: white;
        text-align: center;
        box-shadow: 0 8px 16px rgba(0,0,0,0.15);
        margin-bottom: 25px;
    }
    .main-header h1 { color: #00FF87; font-weight: 800; font-size: 2.8rem; margin-bottom: 10px; }
    .main-header p { font-size: 1.2rem; color: #f1f1f1; margin-bottom: 0; }
</style>
""", unsafe_allow_html=True)

@st.cache_resource
def init_supabase() -> Optional[Client]:
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except Exception:
        return None

supabase = init_supabase()

if "autenticado" not in st.session_state:
    st.session_state.autenticado = False

def verificar_login(username, password):
    if not supabase: return False
    try:
        res = supabase.table("usuarios").select("id, password_hash").eq("username", username).execute()
        if res.data and len(res.data) > 0:
            ok, precisa_migrar = verificar_senha(password, res.data[0]['password_hash'])
            if ok and precisa_migrar:
                try: supabase.table("usuarios").update({"password_hash": gerar_hash_senha(password)}).eq("id", res.data[0]['id']).execute()
                except: pass
            return ok
    except Exception as e: pass
    return False

if not st.session_state.autenticado:
    st.markdown("""
    <style>
        div[data-testid="stAppViewContainer"] { display: flex; align-items: center; }
        .st-key-login_card { max-width: 420px; margin: 4rem auto; padding: 1.5rem 2rem; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 16px rgba(0,0,0,0.12); }
        .login-title { text-align: center; color: #1e3c72; font-weight: 800; font-size: 1.8rem; margin-bottom: 0.3rem; }
    </style>
    """, unsafe_allow_html=True)
    _, col_login, _ = st.columns([1, 1.3, 1])
    with col_login:
        with st.container(border=True, key="login_card"):
            st.markdown('<div class="login-title">⚖️ Autopilot Normativo</div>', unsafe_allow_html=True)
            with st.form("form_login"):
                usuario = st.text_input("Usuário")
                senha = st.text_input("Senha", type="password")
                if st.form_submit_button("Entrar", use_container_width=True):
                    if verificar_login(usuario, senha):
                        st.session_state.autenticado = True
                        st.rerun()
                    else: st.error("❌ Usuário ou senha incorretos.")
    st.stop()

st.markdown("""<div class="main-header"><h1>⚖️ Autopilot Normativo</h1><p>Motor Híbrido OCR e Aprendizado Contínuo com Efeito Cascata</p></div>""", unsafe_allow_html=True)

col_info, col_hist, col_usr, col_logout = st.columns([3, 1.5, 1.5, 1])
with col_info: st.info("💡 **Sistema Autenticado:** Proteção de dados ativa.")
hist_path = "pages/1_Historico.py" if os.path.exists("pages/1_Historico.py") else "pages/historico.py"
usr_path = "pages/usuarios.py"
with col_hist:
    try: st.page_link(hist_path, label="🗄️ Histórico", icon="➡️")
    except: st.markdown(f'<a href="{hist_path.replace("pages/","").replace(".py","")}" target="_top" style="display:block;text-align:center;background:#f0f2f6;border:1px solid #d0d4dc;color:#31333F;padding:0.5rem;border-radius:0.5rem;text-decoration:none;font-weight:500;">➡️ 🗄️ Histórico</a>', unsafe_allow_html=True)
with col_usr:
    try: st.page_link(usr_path, label="👥 Usuários", icon="➡️")
    except: st.markdown(f'<a href="{usr_path.replace("pages/","").replace(".py","")}" target="_top" style="display:block;text-align:center;background:#f0f2f6;border:1px solid #d0d4dc;color:#31333F;padding:0.5rem;border-radius:0.5rem;text-decoration:none;font-weight:500;">➡️ 👥 Usuários</a>', unsafe_allow_html=True)
with col_logout:
    if st.button("Sair", type="secondary", use_container_width=True):
        st.session_state.autenticado = False
        st.rerun()
st.markdown("---")

provedor_escolhido = st.selectbox("🧠 Motor de IA (Hub Multi-IA)", list(PROVEDORES_IA.keys()), key="provedor_ia_select")
modo_processamento = st.radio("⚡ Modo de Processamento", ["Equilibrado", "Rápido", "Máxima Qualidade"], index=0, horizontal=True)

cfg_provedor = PROVEDORES_IA[provedor_escolhido]
api_key = obter_chave_provedor(provedor_escolhido)
if not api_key: api_key = st.text_input(f"Chave da API", type="password")

st.markdown("### 📥 Upload e Tratamento de Atos Normativos")
arquivos_enviados = st.file_uploader("Arraste o documento original ou derivativo (PDF)", type=["pdf"], accept_multiple_files=True, key="uploader_lote")

# =====================================================================
# MOTOR HTML / QUILL
# =====================================================================
class QuillParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.paragraphs = []
        self.current_html = []
        self.stack = []
    def _close_all_tags(self):
        res = ""
        for tags in reversed(self.stack):
            for t in reversed(tags):
                res += f"</{t.split()[0]}>"
        return res
    def _open_all_tags(self):
        res = ""
        for tags in self.stack:
            for t in tags: res += f"<{t}>"
        return res
    def _break_paragraph(self):
        if self.current_html:
            p_text = "".join(self.current_html) + self._close_all_tags()
            if re.sub(r'<[^>]+>', '', p_text).strip(): self.paragraphs.append(p_text.strip())
        self.current_html = []
        if self.stack: self.current_html.append(self._open_all_tags())
    def handle_starttag(self, tag, attrs):
        if tag in ('p', 'br', 'div'): return self._break_paragraph()
        attrs_dict = dict(attrs)
        style = attrs_dict.get('style', '').lower().replace(' ', '')
        cls = attrs_dict.get('class', '').lower()
        added_tags = []
        if tag in ('b', 'strong') or 'font-weight:bold' in style or 'font-weight:700' in style: added_tags.append("b")
        if tag in ('i', 'em') or 'font-style:italic' in style: added_tags.append("i")
        if tag in ('s', 'strike', 'del') or 'text-decoration:line-through' in style or 'ql-strike' in cls: added_tags.append("strike")
        if ('color:rgb(230' in style or 'color:red' in style or 'color:#f00' in style) or (tag == 'font' and attrs_dict.get('color') in ('red', '#f00', '#ff0000')): added_tags.append('font color="red"')
        if added_tags:
            for t in added_tags: self.current_html.append(f"<{t}>")
            self.stack.append(added_tags)
        else: self.stack.append([])
    def handle_endtag(self, tag):
        if tag in ('p', 'br', 'div'): return 
        if self.stack:
            tags_to_close = self.stack.pop()
            for t in reversed(tags_to_close): self.current_html.append(f"</{t.split()[0]}>")
    def handle_startendtag(self, tag, attrs):
        if tag in ('br',): self._break_paragraph()
    def handle_data(self, data):
        if not data: return
        data = data.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\xa0', '&nbsp;')
        self.current_html.append(data)
    def get_paragraphs(self):
        self._break_paragraph()
        return self.paragraphs

def ia_para_editor(texto):
    if not texto: return ""
    texto = texto.replace("<br/>", "</p><p>").replace("<br>", "</p><p>")
    if not texto.startswith("<p>"): texto = f"<p>{texto}</p>"
    texto = re.sub(r'<(strike|del)\b[^>]*>', '<s>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</(strike|del)>', '</s>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<font[^>]*color=[\'"]?(red|#f00|#ff0000|rgb\([^)]+\))[\'"]?[^>]*>', '<span style="color: rgb(230, 0, 0);">', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</font>', '</span>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<b\b[^>]*>', '<strong>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</b>', '</strong>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<i\b[^>]*>', '<em>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</i>', '</em>', texto, flags=re.IGNORECASE)
    return texto.replace("<p></p>", "")

_QUILL_TOOLBAR = [["bold", "italic", "underline", "strike"], [{"color": []}, {"background": []}], [{"list": "ordered"}, {"list": "bullet"}], ["clean"]]

def editor_rico(value, key):
    try: return st_quill(value=value, html=True, toolbar=_QUILL_TOOLBAR, key=key)
    except Exception: return st.text_area("HTML", value=value, key=f"{key}_fallback", label_visibility="collapsed", height=150)

def editor_para_pdf(texto):
    if not texto: return ""
    parser = QuillParser()
    try:
        parser.feed(texto)
        return "<br/>".join(parser.get_paragraphs())
    except Exception: return re.sub(r'</?(span|div|p|ul|li|ol)[^>]*>', '', texto, flags=re.IGNORECASE)

# =====================================================================
# SCHEMAS PYDANTIC E PROMPTS
# =====================================================================
class ArquivoClassificado(BaseModel):
    nome_arquivo_upload: str
    tipo: str = Field(description="'Base' ou 'Alteradora'")
    grupo_id: int
    nome_padronizado_identificado: str
    data_oficial_iso: str
    ato_base_referenciado_tipo: Optional[str] = None
    ato_base_referenciado_numero: Optional[str] = None

class TriagemDocumentos(BaseModel): arquivos: List[ArquivoClassificado]

class MetadadosNorma(BaseModel):
    tipo_documento: str; numero_documento: str; orgao_emissor: str; data_assinatura: str; nome_padronizado: str

class Dispositivo(BaseModel):
    tipo: str; texto_principal_alterada: str; texto_principal_consolidada: str; is_tabela: bool
    tabela_alterada: Optional[List[List[str]]] = None; tabela_consolidada: Optional[List[List[str]]] = None
    texto_pos_tabela_alterada: Optional[str] = None; texto_pos_tabela_consolidada: Optional[str] = None
    nota_remissiva: Optional[str] = Field(default="")

class Consolidacao(BaseModel):
    arquivos_originais_identificados: List[str]; arquivos_alteradores_identificados: List[str]
    norma_base: MetadadosNorma; normas_alteradoras: List[MetadadosNorma]
    cabecalho_complemento: str; orgaos_emissores: str; titulo_portaria: str; ementa: str; preambulo: str
    assinatura_nome: str; assinatura_cargo: str; dispositivos: List[Dispositivo]

SYSTEM_INSTRUCTION_LEGISTECNICA = """
Você é um Especialista Sênior em Técnica Legislativa do Poder Público brasileiro. Regras obrigatórias:
1. FIDELIDADE ABSOLUTA: transcreva o conteúdo de cada dispositivo, preservando formatação (<b>, <i>, <br/>).
2. CRITÉRIO RIGOROSO DE ALTERAÇÃO: Na versão ALTERADA (`texto_principal_alterada`), todo dispositivo alterado/revogado DEVE aparecer com a tag exata: `<strike><font color="red">texto antigo</font></strike>` seguido imediatamente pelo texto novo vigente. Na versão CONSOLIDADA (`texto_principal_consolidada`), exiba apenas a NOVA redação vigente sem riscos.
3. ACÚMULO DE NOTAS (EFEITO CASCATA): Se o texto de origem JÁ POSSUIR uma nota remissiva (ex: '(Redação dada pela PORTARIA X)'), NUNCA APAGUE. Adicione a NOVA nota logo em seguida mantendo a ordem cronológica. O mesmo vale para tabelas.
"""

def _prompt_schema_json(response_schema):
    return "\n\nRESPONDA EXCLUSIVAMENTE COM JSON VÁLIDO QUE OBEDEÇA AO SCHEMA:\n" + json.dumps(response_schema.model_json_schema(), ensure_ascii=False)

def _extrair_json_bruto(texto):
    t = re.sub(r'^```(json)?', '', texto.strip(), flags=re.IGNORECASE).strip()
    t = re.sub(r'```$', '', t.strip()).strip()
    inicio, fim = t.find('{'), t.rfind('}')
    return t[inicio:fim + 1]

def extrair_conteudo_multimodal(file_bytes, nome_arquivo, dpi_ocr=1.5, max_paginas_ocr=None):
    if nome_arquivo.lower().endswith(".docx"): return [f"ARQUIVO DOCX: {nome_arquivo}"]
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        html_text = f"CONTEÚDO DO ARQUIVO {nome_arquivo}:\n\n"
        caracteres_uteis = 0
        for page_num, page in enumerate(doc):
            html_text += f"=== PÁGINA {page_num + 1} ===\n"
            tabelas_bbox = []
            try:
                for tabela in page.find_tables().tables:
                    linhas = tabela.extract()
                    if not linhas: continue
                    caracteres_uteis += sum(len(str(c or "")) for linha in linhas for c in linha)
                    tabelas_bbox.append(fitz.Rect(tabela.bbox))
                    html_text += "[TABELA]\n"
                    for linha in linhas: html_text += " | ".join((str(c).strip() if c is not None else "") for c in linha) + "\n"
                    html_text += "[/TABELA]\n<br/>\n"
            except: pass
            for b in page.get_text("dict", sort=True).get("blocks", []):
                if b.get('type') != 0: continue
                if any(fitz.Rect(b.get("bbox", (0,0,0,0))).intersects(tb) for tb in tabelas_bbox): continue
                bloco_linhas = ""
                for l in b.get("lines", []):
                    linha_span = ""
                    for s in l.get("spans", []):
                        texto = s.get("text", "")
                        if not texto: continue
                        caracteres_uteis += len(texto.strip())
                        flags = s.get("flags", 0)
                        if flags & 2**4: texto = f"<b>{texto}</b>"
                        if flags & 2**1: texto = f"<i>{texto}</i>"
                        linha_span += texto
                    if linha_span.strip(): bloco_linhas += linha_span + " "
                if bloco_linhas.strip(): html_text += bloco_linhas.strip() + "<br/>\n"
            html_text += "<br/>\n"
        if caracteres_uteis < 30 * max(doc.page_count, 1):
            partes = [f"ARQUIVO {nome_arquivo} É UM DOCUMENTO ESCANEADO. Leia o conteúdo visualmente:"]
            for page_num, enumerate_doc in enumerate(doc):
                if max_paginas_ocr and page_num >= max_paginas_ocr: break
                pix = enumerate_doc.get_pixmap(matrix=fitz.Matrix(dpi_ocr, dpi_ocr))
                partes.append({"tipo": "imagem", "mime": "image/jpeg", "dados": pix.tobytes("jpg", jpg_quality=78)})
            return partes
        return [html_text]
    except Exception as e: return [f"Erro: {str(e)}"]

@st.cache_data(show_spinner=False, max_entries=20)
def extrair_conteudo_cache(file_bytes, nome_arquivo, dpi_ocr=1.5, max_paginas_ocr=None):
    return extrair_conteudo_multimodal(file_bytes, nome_arquivo, dpi_ocr, max_paginas_ocr)

def _chamar_motor(motor, chave, itens, schema, thinking_level, modelos):
    if motor == "gemini":
        client = genai.Client(api_key=chave)
        config = types.GenerateContentConfig(response_mime_type="application/json", response_schema=schema, system_instruction=SYSTEM_INSTRUCTION_LEGISTECNICA, thinking_config=types.ThinkingConfig(thinking_level=thinking_level))
        partes = [types.Part.from_bytes(data=it["dados"], mime_type=it["mime"]) if isinstance(it, dict) else it for it in itens]
        for mod in modelos:
            for _ in range(3):
                try:
                    resp = client.models.generate_content(model=mod, contents=partes, config=config)
                    return schema.model_validate(json.loads(resp.text))
                except Exception as e: 
                    if "429" in str(e) or "503" in str(e): time.sleep(5); continue
                    break
        raise Exception("Gemini falhou.")
    else:
        textos, imagens = [], []
        for it in itens:
            if isinstance(it, dict): imagens.append(it)
            else: textos.append(it)
        texto_final = "\n\n".join(textos) + _prompt_schema_json(schema)
        conteudo = [{"type": "text", "text": texto_final}]
        for img in imagens: conteudo.append({"type": "image_url", "image_url": {"url": f"data:{img['mime']};base64,{base64.b64encode(img['dados']).decode()}"}})
        mensagens = [{"role": "system", "content": SYSTEM_INSTRUCTION_LEGISTECNICA}, {"role": "user", "content": conteudo}]
        
        client = None
        if motor == "groq" and Groq: client = Groq(api_key=chave)
        elif motor == "openrouter" and OpenAI: client = OpenAI(api_key=chave, base_url="https://openrouter.ai/api/v1")
        elif motor == "mistral" and Mistral: client = Mistral(api_key=chave)
        if not client: raise Exception(f"Cliente {motor} não disponível.")

        for mod in modelos:
            for _ in range(3):
                try:
                    fn = client.chat.complete if motor == "mistral" else client.chat.completions.create
                    resp = fn(model=mod, messages=mensagens, response_format={"type": "json_object"}, temperature=0.2)
                    return schema.model_validate(json.loads(_extrair_json_bruto(resp.choices[0].message.content)))
                except Exception as e:
                    if "429" in str(e) or "503" in str(e): time.sleep(5); continue
                    break
        raise Exception(f"{motor} falhou.")

def executar_com_fallback(chave, itens, schema, provedor, thinking_level="high"):
    cfg = PROVEDORES_IA[provedor]
    try: res = _chamar_motor(cfg["motor"], chave, itens, schema, thinking_level, cfg["modelos"])
    except Exception as e:
        for p_alt, cfg_alt in PROVEDORES_IA.items():
            if p_alt == provedor: continue
            k_alt = obter_chave_provedor(p_alt)
            if not k_alt: continue
            try: res = _chamar_motor(cfg_alt["motor"], k_alt, itens, schema, thinking_level, cfg_alt["modelos"]); break
            except: pass
        else: raise e
    class _Resp: text = res.model_dump_json()
    return _Resp()

def converter_para_iso(data_str):
    if not data_str: return None
    data_str = data_str.strip()
    if re.match(r'^\d{4}-\d{2}-\d{2}$', data_str): return data_str
    try: return datetime.strptime(data_str, "%d/%m/%Y").strftime("%Y-%m-%d")
    except: return None

def corrigir_posicionamento_tabela(consolidacao: dict):
    if not isinstance(consolidacao, dict): return consolidacao
    for disp in consolidacao.get("dispositivos", []):
        if not disp.get("is_tabela"): continue
        txt_alt, txt_pos_alt = disp.get("texto_principal_alterada") or "", disp.get("texto_pos_tabela_alterada") or ""
        if "redação dada pelo" in txt_pos_alt.lower() or "nova redação" in txt_pos_alt.lower(): continue
        partes = re.split(r'<br\s*/?>\s*<br\s*/?>', txt_alt, flags=re.IGNORECASE)
        nova_redacao = None
        if len(partes) >= 2:
            if '<strike' not in partes[-1].lower() and '<s>' not in partes[-1].lower():
                nova_redacao = partes[-1].strip()
                texto_antigo = "<br/><br/>".join(partes[:-1]).strip()
                disp["texto_principal_alterada"] = texto_antigo + "<br/><br/>"
                disp["texto_pos_tabela_alterada"] = nova_redacao + ("<br/><br/>" + txt_pos_alt if txt_pos_alt else "")
    return consolidacao

def resgatar_memoria():
    memoria = ""
    if supabase:
        try:
            res = supabase.table("memoria_de_correcoes").select("*").order("id", desc=True).limit(5).execute()
            if res.data:
                memoria = "\n\n⚠️ HISTÓRICO DE CORREÇÕES (Não repita os erros da IA):\n"
                for m in res.data: memoria += f"- Erro: {m['texto_ia']}\n- Correção: {m['texto_corrigido']}\n\n"
        except: pass
    return memoria

# =====================================================================
# LÓGICA DE DUAS ETAPAS (LEITURA INICIAL E CONSOLIDAÇÃO)
# =====================================================================
def ler_arquivos_iniciais(arquivos, api_key, provedor, thinking_level, dpi_ocr, max_paginas_ocr, progresso=None):
    textos_extraidos = {}
    max_workers = 2 if thinking_level == "low" else 4
    with ThreadPoolExecutor(max_workers=min(max_workers, max(1, len(arquivos)))) as ex:
        futuros = {submit_com_contexto(ex, extrair_conteudo_cache, arq.getvalue(), arq.name, dpi_ocr, max_paginas_ocr): arq.name for arq in arquivos}
        for idx, fut in enumerate(as_completed(futuros)):
            textos_extraidos[futuros[fut]] = fut.result()
            if progresso: progresso.progress((idx + 1) / (len(arquivos) * 2), text=f"Extraindo conteúdo de {futuros[fut]}...")

    prompt_triagem = ["Analise e identifique os documentos (Base ou Alteradora) e as referências cruzadas. ARQUIVOS:"]
    for k, v in textos_extraidos.items(): prompt_triagem.extend(v if isinstance(v, list) else [v])
    
    if progresso: progresso.progress(0.5, text="Classificando documentos...")
    resp_triagem = executar_com_fallback(api_key, prompt_triagem, TriagemDocumentos, provedor, "low")
    triagem_dict = {a['nome_arquivo_upload']: a for a in json.loads(resp_triagem.text).get("arquivos", [])}

    atos_iniciais = []
    total = len(arquivos)
    for idx, f in enumerate(arquivos):
        if progresso: progresso.progress(0.5 + 0.5 * ((idx + 1) / total), text=f"Estruturando {f.name}...")
        txt_partes = textos_extraidos[f.name]
        prompt = ["DOCUMENTO LIDO:\n"] + (txt_partes if isinstance(txt_partes, list) else [txt_partes])
        prompt.append("\n\nExtraia e estruture este ato normativo individualmente no formato JSON. "
                      "ATENÇÃO: Como este é um processamento de LEITURA INICIAL (etapa 1), preencha APENAS o campo "
                      "`texto_principal_consolidada` (e `tabela_consolidada` se houver) de cada dispositivo com o texto "
                      "original lido. Deixe todos os campos referentes a 'alterada' VAZIOS. Não faça cruzamentos.")
        resp = executar_com_fallback(api_key, prompt, Consolidacao, provedor, thinking_level)
        ato_json = json.loads(resp.text)
        ato_json['_upload_name'] = f.name
        ato_json['_triagem'] = triagem_dict.get(f.name, {"tipo": "Base", "nome_padronizado_identificado": f.name})
        ato_json['fase'] = 'leitura'
        atos_iniciais.append(ato_json)
        
    return atos_iniciais

def consolidar_documentos(ato_alterador, base_id_db, api_key, provedor, thinking_level):
    base_data = supabase.table("portarias_base").select("documento_consolidado_json").eq("id", base_id_db).execute().data
    if not base_data: raise Exception("Ato base não encontrado no banco.")
    base_json = base_data[0]['documento_consolidado_json']
    
    memoria = resgatar_memoria()
    prompt = [
        f"ESTADO ATUAL DO BANCO (Ato Base Vigente e com alterações anteriores aplicadas):\n{json.dumps(base_json)}\n\n",
        f"NOVO ATO ALTERADOR (Documento lido agora):\n{json.dumps(ato_alterador)}\n\n",
        "Sua tarefa é cruzar os dados gerando o novo estado Consolidado final.\n"
        "Siga a marcação <strike><font color='red'>... na versão Alterada e deixe o texto limpo na Consolidada.\n"
        "NUNCA APAGUE as notas remissivas que já existem na Base; apenas acumule as novas notas no final da linha (Efeito Cascata).\n"
        + memoria
    ]
    resp = executar_com_fallback(api_key, prompt, Consolidacao, provedor, thinking_level)
    resultado = json.loads(resp.text)
    resultado = corrigir_posicionamento_tabela(resultado)
    return resultado

def gerar_html_dinamico(consolidacao_dict, tipo_versao):
    titulo_doc = f"Versão {'Alterada' if tipo_versao=='alterada' else 'Consolidada'}"
    html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{titulo_doc}</title><style>@page {{ size: A4; margin: 2.5cm 2cm; @bottom-center {{ content: 'Nota: Este documento possui caráter estritamente consultivo e informativo, não substituindo o texto original publicado no Diário Oficial.'; font-size: 9pt; font-style: italic; color: #333; }} }} body {{ font-family: 'Times New Roman', serif; font-size: 11pt; line-height: 1.5; text-align: justify; }} .topo, .titulo, .orgaos, .capitulo, .assinatura {{ text-align: center; font-weight: bold; }} .ementa {{ margin-left: 45%; }} .dispositivo {{ text-indent: 40px; }} table {{ width: 100%; border-collapse: collapse; }} td, th {{ border: 1px solid black; padding: 6px; }} strike, font[color='red'] {{ color: red !important; text-decoration: line-through; }} </style></head><body>"
    html += f"<div class='orgaos'>{consolidacao_dict.get('orgaos_emissores','')}</div><div class='titulo'>{consolidacao_dict.get('titulo_portaria','')}</div><div class='ementa'>{consolidacao_dict.get('ementa','')}</div><div class='preambulo'>{consolidacao_dict.get('preambulo','')}</div>"
    for d in consolidacao_dict.get('dispositivos', []):
        t_prin = d.get(f'texto_principal_{tipo_versao}', '')
        if t_prin: html += f"<div class='dispositivo'>{t_prin}</div>"
        if d.get('is_tabela') and d.get(f'tabela_{tipo_versao}'):
            html += "<table>"
            for linha in d.get(f'tabela_{tipo_versao}'): html += "<tr>" + "".join([f"<td>{cel}</td>" for cel in linha]) + "</tr>"
            html += "</table>"
            t_pos = d.get(f'texto_pos_tabela_{tipo_versao}', '')
            if t_pos: html += f"<div class='dispositivo'>{t_pos}</div>"
    html += f"<div class='assinatura'>{consolidacao_dict.get('assinatura_nome','')}<br>{consolidacao_dict.get('assinatura_cargo','')}</div></body></html>"
    return html

def gerar_pdf_dinamico(consolidacao_dict, tipo_versao):
    if not HAS_WEASYPRINT: raise Exception("WeasyPrint indisponível")
    html_str = gerar_html_dinamico(consolidacao_dict, tipo_versao)
    buffer = io.BytesIO()
    WeasyHTML(string=html_str).write_pdf(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_docx_dinamico(consolidacao_dict, tipo_versao):
    doc = docx.Document()
    doc.add_paragraph(consolidacao_dict.get('titulo_portaria', '')).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for d in consolidacao_dict.get('dispositivos', []):
        doc.add_paragraph(re.sub(r'<[^>]+>', '', d.get(f'texto_principal_{tipo_versao}', ''))).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    b = io.BytesIO(); doc.save(b); b.seek(0)
    return b.getvalue()

# =====================================================================
# UI E FLUXO PRINCIPAL
# =====================================================================
if "atos_estruturados" not in st.session_state: st.session_state.atos_estruturados = []
if "lote_processado_id" not in st.session_state: st.session_state.lote_processado_id = None

current_lote_id = hashlib.md5(str([f.name for f in arquivos_enviados] if arquivos_enviados else "").encode()).hexdigest()

if arquivos_enviados and current_lote_id != st.session_state.lote_processado_id:
    st.session_state.atos_estruturados = []
    st.session_state.lote_processado_id = current_lote_id

if arquivos_enviados and not st.session_state.atos_estruturados:
    if st.button("1. Ler e Estruturar Documento(s)", type="primary"):
        with st.spinner("Lendo PDFs e Estruturando (Etapa 1/2)..."):
            dpi = 1.2 if modo_processamento == "Rápido" else 1.5
            max_p = 10 if modo_processamento == "Rápido" else (20 if modo_processamento == "Equilibrado" else None)
            tl = "low" if modo_processamento == "Rápido" else ("medium" if modo_processamento == "Equilibrado" else "high")
            prog = st.progress(0.0)
            
            try:
                st.session_state.atos_estruturados = ler_arquivos_iniciais(arquivos_enviados, api_key, provedor_escolhido, tl, dpi, max_p, prog)
            except Exception as e:
                st.error(f"Erro ao ler arquivos: {e}")
            finally:
                prog.empty()
        st.rerun()

if st.session_state.atos_estruturados:
    st.markdown("---")
    st.markdown("### 📝 Análise e Consolidação")
    
    for i, ato in enumerate(st.session_state.atos_estruturados):
        fase = ato.get('fase', 'leitura')
        triagem = ato.get('_triagem', {})
        tipo_ato = triagem.get('tipo', 'Base')
        nome_base_ou_alt = triagem.get('nome_padronizado_identificado', ato.get('_upload_name', 'Documento'))

        mostrar_expander = fase not in ['concluido', 'pendente']
        
        with st.expander(f"Fase: {fase.upper()} → **{ato.get('_upload_name')}** ({tipo_ato})", expanded=mostrar_expander):
            
            if fase == 'leitura':
                ato['titulo_portaria'] = st.text_input("Título do Ato Normativo", ato.get('titulo_portaria', ''), key=f"t_{i}")
                
                st.markdown("**Ementa**")
                ato['ementa'] = editor_para_pdf(editor_rico(ia_para_editor(ato.get('ementa', '')), f"q_ementa_{i}"))
                
                st.markdown("**Preâmbulo e Considerandos**")
                ato['preambulo'] = editor_para_pdf(editor_rico(ia_para_editor(ato.get('preambulo', '')), f"q_preambulo_{i}"))
                
                st.markdown("#### Dispositivos (Artigos, Parágrafos, Incisos)")
                for j, disp in enumerate(ato.get("dispositivos", [])):
                    st.markdown(f"**{disp.get('tipo', 'Dispositivo').upper()} {j+1}**")
                    val_orig = ia_para_editor(disp.get('texto_principal_consolidada', ''))
                    disp['texto_principal_consolidada'] = editor_para_pdf(editor_rico(val_orig, f"orig_{i}_{j}"))
                    
                    if disp.get('is_tabela'):
                        disp['tabela_consolidada'] = st.data_editor(disp.get('tabela_consolidada') or [[""]], key=f"tbc_orig_{i}_{j}")
                        disp['texto_pos_tabela_consolidada'] = st.text_area("Pós-tabela", value=disp.get('texto_pos_tabela_consolidada', ''), key=f"tpc_orig_{i}_{j}")
                
                st.markdown("---")

                if tipo_ato == "Base":
                    ja_existe_base = False
                    if supabase:
                        check_base = supabase.table("portarias_base").select("id").eq("nome_padronizado", nome_base_ou_alt).execute()
                        ja_existe_base = bool(check_base.data)

                    if ja_existe_base:
                        st.warning(f"⚠️ Atenção: O Ato Base '{nome_base_ou_alt}' já se encontra cadastrado no banco de dados.")
                        btn_label = "🔄 Substituir Ato Base no Banco"
                        btn_type = "secondary"
                    else:
                        btn_label = "💾 Salvar este Ato Base no Banco de Dados"
                        btn_type = "primary"

                    if st.button(btn_label, key=f"btn_salv_base_{i}", type=btn_type):
                        if supabase:
                            dados_db = {
                                "tipo_documento": ato.get('norma_base', {}).get('tipo_documento', 'Norma'),
                                "numero_documento": ato.get('norma_base', {}).get('numero_documento', 'S/N'),
                                "orgao_emissor": ato.get('norma_base', {}).get('orgao_emissor', ''),
                                "data_assinatura": converter_para_iso(ato.get('norma_base', {}).get('data_assinatura')),
                                "nome_padronizado": nome_base_ou_alt,
                                "titulo_original": ato.get("titulo_portaria"),
                                "documento_consolidado_json": ato
                            }
                            try:
                                supabase.table("portarias_base").upsert(dados_db, on_conflict="nome_padronizado").execute()
                                if ja_existe_base:
                                    st.success(f"O Ato Base '{nome_base_ou_alt}' foi substituído com sucesso!")
                                else:
                                    st.success(f"O Ato Base '{nome_base_ou_alt}' foi cadastrado no banco com sucesso e já está disponível para receber derivações!")
                                ato['fase'] = 'concluido'
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erro ao salvar: {e}")
                        else:
                            st.error("Sem conexão com o DB.")
                            
                elif tipo_ato == "Alteradora":
                    ref_t = triagem.get('ato_base_referenciado_tipo', 'Desconhecido')
                    ref_n = str(triagem.get('ato_base_referenciado_numero', '')).strip()
                    
                    ja_existe_alt = False
                    if supabase:
                        check_alt = supabase.table("portarias_alteradoras").select("id").eq("nome_padronizado", nome_base_ou_alt).execute()
                        ja_existe_alt = bool(check_alt.data)
                    
                    if ja_existe_alt:
                        st.warning(f"⚠️ Atenção: A portaria alteradora '{nome_base_ou_alt}' já foi processada e vinculada a um ato no banco de dados.")
                        btn_txt_check = "🔄 Substituir/Reprocessar Vínculo"
                        btn_type_check = "secondary"
                    else:
                        st.info(f"🔗 O sistema identificou que este ato deriva da norma: **{ref_t} {ref_n}**")
                        btn_txt_check = "💾 Buscar Relações no Banco"
                        btn_type_check = "primary"

                    if st.button(btn_txt_check, key=f"btn_check_{i}", type=btn_type_check):
                        if not supabase:
                            st.error("Sem conexão com o Banco de Dados.")
                        else:
                            bases_db = []
                            if ref_n and ref_n != 'Desconhecido':
                                bases_db = supabase.table("portarias_base").select("id, nome_padronizado, data_assinatura").ilike("numero_documento", f"%{ref_n}%").execute().data or []
                                if not bases_db:
                                    bases_db = supabase.table("portarias_base").select("id, nome_padronizado, data_assinatura").ilike("nome_padronizado", f"%{ref_n}%").execute().data or []
                            
                            if bases_db:
                                ato['candidatos_base'] = bases_db
                                ato['fase'] = 'vincular'
                                ato['ja_existe_alt'] = ja_existe_alt
                                st.rerun()
                            else:
                                try:
                                    supabase.table("atos_importados").insert({
                                        "nome_arquivo_original": ato.get('_upload_name', ''),
                                        "texto_integra": json.dumps(ato),
                                        "tipo_documento": ato.get('norma_base', {}).get('tipo_documento', None),
                                        "numero_documento": ato.get('norma_base', {}).get('numero_documento', None),
                                        "data_assinatura": converter_para_iso(ato.get('norma_base', {}).get('data_assinatura')),
                                        "orgao_emissor": ato.get('norma_base', {}).get('orgao_emissor', None),
                                        "ato_base_referenciado_tipo": ref_t,
                                        "ato_base_referenciado_numero": ref_n,
                                        "status": "pendente",
                                        "estrutura_json": ato
                                    }).execute()
                                    ato['fase'] = 'pendente'
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Erro ao salvar pendência: {e}")

            elif fase == 'vincular':
                ref_t = triagem.get('ato_base_referenciado_tipo', 'Desconhecido')
                ref_n = triagem.get('ato_base_referenciado_numero', 'Desconhecido')
                st.info(f"🔗 O sistema identificou que este ato deriva da norma: **{ref_t} {ref_n}** e localizou correspondências no banco.")
                
                candidatos = ato.get('candidatos_base', [])
                opcoes = {b['id']: f"{b['nome_padronizado']} (Data: {b.get('data_assinatura', 'S/D')})" for b in candidatos}
                sel_base = st.selectbox("Selecione o Ato Base correto para vincular e consolidar:", options=list(opcoes.keys()), format_func=lambda x: opcoes[x], key=f"sel_base_{i}")
                
                ja_existe_alt = ato.get('ja_existe_alt', False)
                if ja_existe_alt:
                    btn_txt_consol = "🔄 Analisar, Substituir e Consolidar"
                    btn_type_consol = "secondary"
                else:
                    btn_txt_consol = "🔍 Gerar Versões Alterada e Consolidada"
                    btn_type_consol = "primary"
                
                if st.button(btn_txt_consol, key=f"btn_consol_{i}", type=btn_type_consol):
                    with st.spinner("Buscando histórico no banco e aplicando Efeito Cascata..."):
                        tl = "low" if modo_processamento == "Rápido" else ("medium" if modo_processamento == "Equilibrado" else "high")
                        try:
                            novo_ato = consolidar_documentos(ato, sel_base, api_key, provedor_escolhido, tl)
                            novo_ato['_triagem'] = triagem
                            novo_ato['_id_base_vinculada'] = sel_base
                            novo_ato['_upload_name'] = ato.get('_upload_name')
                            novo_ato['fase'] = 'consolidacao'
                            st.session_state.atos_estruturados[i] = novo_ato
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro na análise de consolidação: {e}")

            elif fase == 'consolidacao':
                ato['titulo_portaria'] = st.text_input("Título do Ato Normativo", ato.get('titulo_portaria', ''), key=f"t_cons_{i}")
                
                st.markdown("**Ementa**")
                ato['ementa'] = editor_para_pdf(editor_rico(ia_para_editor(ato.get('ementa', '')), f"q_ementa_cons_{i}"))
                
                st.markdown("**Preâmbulo e Considerandos**")
                ato['preambulo'] = editor_para_pdf(editor_rico(ia_para_editor(ato.get('preambulo', '')), f"q_preambulo_cons_{i}"))
                
                st.markdown("#### Dispositivos (Artigos, Parágrafos, Incisos)")
                for j, disp in enumerate(ato.get("dispositivos", [])):
                    st.markdown(f"**{disp.get('tipo', 'Dispositivo').upper()} {j+1}**")
                    c_alt, c_cons = st.columns(2)
                    
                    with c_alt:
                        st.markdown("*Versão Alterada*")
                        val_alt = ia_para_editor(disp.get('texto_principal_alterada', ''))
                        alt_editada = editor_rico(value=val_alt, key=f"q_alt_{i}_{j}")
                        if alt_editada is not None: disp['texto_principal_alterada'] = editor_para_pdf(alt_editada)
                        
                    with c_cons:
                        st.markdown("*Versão Consolidada*")
                        val_cons = ia_para_editor(disp.get('texto_principal_consolidada', ''))
                        cons_editada = editor_rico(value=val_cons, key=f"q_cons_{i}_{j}")
                        if cons_editada is not None: disp['texto_principal_consolidada'] = editor_para_pdf(cons_editada)
                    
                    disp['nota_remissiva'] = st.text_input("Nota Remissiva (Indexação)", value=disp.get('nota_remissiva', ''), key=f"nota_{i}_{j}", label_visibility="collapsed")
                    st.markdown("---")

                    if disp.get('is_tabela'):
                        st.markdown("*Tabela / Anexo — revise linha a linha*")
                        t_alt, t_cons = st.columns(2)
                        with t_alt:
                            tab_alt_edit = st.data_editor(disp.get('tabela_alterada') or [[""]], key=f"tab_alt_{i}_{j}", num_rows="dynamic", use_container_width=True)
                            disp['tabela_alterada'] = tab_alt_edit if isinstance(tab_alt_edit, list) else disp.get('tabela_alterada')
                            pos_alt = st.text_area("Texto após a tabela (Alterada)", value=disp.get('texto_pos_tabela_alterada') or "", key=f"pos_alt_{i}_{j}")
                            disp['texto_pos_tabela_alterada'] = pos_alt
                        with t_cons:
                            tab_cons_edit = st.data_editor(disp.get('tabela_consolidada') or [[""]], key=f"tab_cons_{i}_{j}", num_rows="dynamic", use_container_width=True)
                            disp['tabela_consolidada'] = tab_cons_edit if isinstance(tab_cons_edit, list) else disp.get('tabela_consolidada')
                            pos_cons = st.text_area("Texto após a tabela (Consolidada)", value=disp.get('texto_pos_tabela_consolidada') or "", key=f"pos_cons_{i}_{j}")
                            disp['texto_pos_tabela_consolidada'] = pos_cons

                st.markdown("### 📥 Opções de Banco e Exportação")
                
                if st.button("💾 Confirmar e Salvar Versões no Banco", key=f"btn_salvar_versoes_{i}", type="primary"):
                    if supabase:
                        try:
                            supabase.table("portarias_base").update({"documento_consolidado_json": ato}).eq("id", ato['_id_base_vinculada']).execute()
                            desc = f"Alteração processada via: {ato.get('_upload_name')}"
                            alt_aplicadas = []
                            
                            for n_alt in ato.get('normas_alteradoras', []):
                                alt_nome = n_alt.get('nome_padronizado')
                                alt_aplicadas.append(alt_nome)
                                check_alt = supabase.table("portarias_alteradoras").select("id").eq("portaria_base_id", ato['_id_base_vinculada']).eq("nome_padronizado", alt_nome).execute()
                                if not check_alt.data:
                                    supabase.table("portarias_alteradoras").insert({
                                        "portaria_base_id": ato['_id_base_vinculada'],
                                        "tipo_documento": n_alt.get('tipo_documento', 'Alteradora'),
                                        "numero_documento": n_alt.get('numero_documento', 'S/N'),
                                        "orgao_emissor": n_alt.get('orgao_emissor', ''),
                                        "data_assinatura": converter_para_iso(n_alt.get('data_assinatura')),
                                        "nome_padronizado": alt_nome,
                                        "arquivo_nome_original": ato.get('_upload_name', '')
                                    }).execute()
                            
                            supabase.table("versoes_documentos").insert({
                                "portaria_base_id": ato['_id_base_vinculada'],
                                "tipo_versao": "alterada",
                                "estado_json": ato,
                                "alteradoras_aplicadas": alt_aplicadas,
                                "descricao": desc
                            }).execute()
                            
                            supabase.table("versoes_documentos").insert({
                                "portaria_base_id": ato['_id_base_vinculada'],
                                "tipo_versao": "consolidada",
                                "estado_json": ato,
                                "alteradoras_aplicadas": alt_aplicadas,
                                "descricao": desc
                            }).execute()
                            
                            ato['fase'] = 'concluido'
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao salvar versões no histórico: {e}")

                c_html, c_pdf, c_docx = st.columns(3)
                arq_base = nome_base_ou_alt.replace(' ', '_').replace('/', '-')
                
                try:
                    c_html.download_button("🌐 Baixar HTML (Alt)", gerar_html_dinamico(ato, "alterada"), f"{arq_base}_Alt.html", "text/html", key=f"ha_cons_{i}")
                    c_html.download_button("🌐 Baixar HTML (Cons)", gerar_html_dinamico(ato, "consolidada"), f"{arq_base}_Cons.html", "text/html", key=f"hc_cons_{i}")
                except: pass
                try:
                    c_pdf.download_button("📄 Baixar PDF (Alt)", gerar_pdf_dinamico(ato, "alterada"), f"{arq_base}_Alt.pdf", "application/pdf", key=f"pa_cons_{i}")
                    c_pdf.download_button("📄 Baixar PDF (Cons)", gerar_pdf_dinamico(ato, "consolidada"), f"{arq_base}_Cons.pdf", "application/pdf", key=f"pc_cons_{i}")
                except: pass
                try:
                    c_docx.download_button("📝 Baixar DOCX (Alt)", gerar_docx_dinamico(ato, "alterada"), f"{arq_base}_Alt.docx", "application/vnd.openxmlformats", key=f"da_cons_{i}")
                    c_docx.download_button("📝 Baixar DOCX (Cons)", gerar_docx_dinamico(ato, "consolidada"), f"{arq_base}_Cons.docx", "application/vnd.openxmlformats", key=f"dc_cons_{i}")
                except: pass

            elif fase == 'concluido':
                if tipo_ato == "Base":
                    st.success("✅ O Ato Base foi cadastrado/atualizado com sucesso e está disponível para derivações.")
                else:
                    st.success("✅ Análise concluída! O banco de dados e as versões salvas foram atualizados.")
                
            elif fase == 'pendente':
                ref_t = triagem.get('ato_base_referenciado_tipo', 'Desconhecido')
                ref_n = triagem.get('ato_base_referenciado_numero', 'Desconhecido')
                st.warning(f"⚠️ O ato base relacionado ({ref_t} {ref_n}) não foi encontrado no banco de dados.")
                st.success("✅ Este ato foi armazenado na tabela de Pendências. Você poderá consolidá-lo via Histórico quando o ato original for enviado.")

    st.markdown("---")
    if st.button("🔄 Limpar Tela e Fazer Novo Upload", type="secondary"): 
        st.session_state.atos_estruturados = []
        st.rerun()
