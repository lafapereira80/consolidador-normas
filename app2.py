import streamlit as st
import tempfile
import io
import json
import os
import re
import time
import copy
import hashlib
import base64
from html.parser import HTMLParser
from html import unescape
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from streamlit.runtime.scriptrunner import add_script_run_ctx, get_script_run_ctx
from google import genai
from google.genai import types
from pydantic import BaseModel, Field, ValidationError
from typing import List, Optional

# Importação para Supabase, Word, Leitura de PDF e Editor Web
from supabase import create_client, Client
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from auth_utils import gerar_hash_senha, verificar_senha
from streamlit_quill import st_quill

try:
    from groq import Groq
except ImportError:
    Groq = None
try:
    from mistralai import Mistral
except ImportError:
    try:
        from mistralai.client import Mistral
    except ImportError:
        Mistral = None
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

# ----------------- HUB MULTI-IA: PROVEDORES E MODELOS MAIS RECENTES -----------------
# Cada provedor tem uma cadeia de fallback (mais novo/capaz -> mais estável) dentro
# dele mesmo. A troca de PROVEDOR é feita pelo usuário na tela.
PROVEDORES_IA = {
    "Google Gemini": {
        "motor": "gemini",
        "modelos": ["gemini-3.7-flash", "gemini-3.6-flash", "gemini-3.5-flash"],
        "secret": "GEMINI_API_KEY",
    },
    "Groq (Llama / GPT-OSS)": {
        "motor": "groq",
        "modelos": ["openai/gpt-oss-120b", "openai/gpt-oss-20b", "qwen/qwen3.6-27b"],
        "secret": "GROQ_API_KEY",
    },
    "OpenRouter (Qwen / DeepSeek / Llama)": {
        "motor": "openrouter",
        "modelos": ["deepseek/deepseek-v4-flash", "qwen/qwen3.5-plus", "meta-llama/llama-4-maverick"],
        "secret": "OPENROUTER_API_KEY",
    },
    "Mistral AI (Small / Nemo)": {
        "motor": "mistral",
        "modelos": ["mistral-small-latest", "open-mistral-nemo"],
        "secret": "MISTRAL_API_KEY",
    },
}

def submit_com_contexto(executor, fn, *args, **kwargs):
    """Envia uma tarefa ao ThreadPoolExecutor propagando o ScriptRunContext do
    Streamlit — sem isso, st.toast/st.info/st.warning chamados de dentro da
    thread são descartados silenciosamente (a sessão do usuário não é
    identificada), fazendo avisos de retry/erro desaparecerem sob carga."""
    ctx = get_script_run_ctx()
    def _wrapper(*a, **kw):
        if ctx is not None:
            add_script_run_ctx(threading.current_thread(), ctx)
        return fn(*a, **kw)
    return executor.submit(_wrapper, *args, **kwargs)

def obter_chave_provedor(nome_provedor):
    cfg = PROVEDORES_IA[nome_provedor]
    try:
        return st.secrets[cfg["secret"]]
    except Exception:
        try:
            return st.secrets["api_keys"][cfg["secret"]]
        except Exception:
            return os.environ.get(cfg["secret"], "")

# Tenta carregar o WeasyPrint (Gerador de PDF Avançado)
try:
    from weasyprint import HTML as WeasyHTML, CSS as WeasyCSS
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False

# ----------------- CONFIGURAÇÃO DA PÁGINA -----------------
st.set_page_config(page_title="Autopilot Normativo", page_icon="⚖️", layout="wide", initial_sidebar_state="collapsed")

# ----------------- BLOQUEIO TOTAL DO MENU LATERAL E CSS GLOBAL -----------------
st.markdown("""
<style>
    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="collapsedControl"] { display: none !important; }
    
    .block-container { padding-top: 2rem; }
    .main-header {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        padding: 30px 20px;
        border-radius: 12px;
        color: white;
        text-align: center;
        box-shadow: 0 8px 16px rgba(0,0,0,0.15);
        margin-bottom: 25px;
    }
    .main-header h1 { color: #00FF87; font-weight: 800; font-size: 2.8rem; margin-bottom: 10px; }
    .main-header p { font-size: 1.2rem; color: #f1f1f1; margin-bottom: 0; }
</style>
""", unsafe_allow_html=True)

# ----------------- CONEXÃO COM SUPABASE -----------------
@st.cache_resource
def init_supabase() -> Optional[Client]:
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except Exception:
        return None

supabase = init_supabase()

# ----------------- SISTEMA DE AUTENTICAÇÃO (LOGIN) -----------------
if "autenticado" not in st.session_state:
    st.session_state.autenticado = False

def verificar_login(username, password):
    if not supabase:
        st.error("⚠️ Erro de conexão com o Banco de Dados.")
        return False
    try:
        res = supabase.table("usuarios").select("id, password_hash").eq("username", username).execute()
        if res.data and len(res.data) > 0:
            ok, precisa_migrar = verificar_senha(password, res.data[0]['password_hash'])
            if ok and precisa_migrar:
                try:
                    supabase.table("usuarios").update({"password_hash": gerar_hash_senha(password)}).eq("id", res.data[0]['id']).execute()
                except Exception:
                    pass
            return ok
    except Exception as e:
        st.error(f"Erro ao verificar credenciais: {e}")
    return False

if not st.session_state.autenticado:
    st.markdown("""
    <style>
        div[data-testid="stAppViewContainer"] { display: flex; align-items: center; }
        .st-key-login_card {
            max-width: 420px;
            margin: 4rem auto;
            padding: 1.5rem 2rem;
            background: #ffffff;
            border-radius: 12px;
            box-shadow: 0 4px 16px rgba(0,0,0,0.12);
        }
        .st-key-login_card [data-testid="stForm"] { border: none; padding: 0; }
        .login-title { text-align: center; color: #1e3c72; font-weight: 800; font-size: 1.8rem; margin-bottom: 0.3rem; }
        .login-subtitle { text-align: center; color: #666; margin-bottom: 1.5rem; }
    </style>
    """, unsafe_allow_html=True)

    _, col_login, _ = st.columns([1, 1.3, 1])
    with col_login:
        with st.container(border=True, key="login_card"):
            st.markdown('<div class="login-title">⚖️ Autopilot Normativo</div>', unsafe_allow_html=True)
            st.markdown('<div class="login-subtitle">Acesso Restrito ao Sistema</div>', unsafe_allow_html=True)

            with st.form("form_login"):
                usuario = st.text_input("Usuário", placeholder="Digite seu usuário")
                senha = st.text_input("Senha", type="password", placeholder="Digite sua senha")
                btn_login = st.form_submit_button("Entrar no Sistema", use_container_width=True)

                if btn_login:
                    if verificar_login(usuario, senha):
                        st.session_state.autenticado = True
                        st.rerun()
                    else:
                        st.error("❌ Usuário ou senha incorretos.")
    st.stop()

# =====================================================================
# ÁREA AUTENTICADA DO SISTEMA
# =====================================================================

st.markdown("""
<div class="main-header">
    <h1>⚖️ Autopilot Normativo</h1>
    <p>Motor Híbrido OCR com Editor Visual e Aprendizado Contínuo</p>
</div>
""", unsafe_allow_html=True)

col_info, col_hist, col_usr, col_logout = st.columns([3, 1.5, 1.5, 1])

with col_info:
    st.info("💡 **Sistema Autenticado:** Proteção de dados ativa.")

hist_path = "pages/historico.py"
usr_path = "pages/usuarios.py"

if os.path.exists("pages"):
    for f in os.listdir("pages"):
        if "historico" in f.lower() and f.endswith(".py"): hist_path = f"pages/{f}"
        if "usuario" in f.lower() and f.endswith(".py"): usr_path = f"pages/{f}"

with col_hist:
    try:
        st.page_link(hist_path, label="🗄️ Histórico", icon="➡️")
    except:
        nome_pagina = hist_path.replace("pages/", "").replace(".py", "")
        st.markdown(f'<a href="{nome_pagina}" target="_top" style="display: block; text-align: center; background-color: #f0f2f6; border: 1px solid #d0d4dc; color: #31333F !important; padding: 0.5rem; border-radius: 0.5rem; text-decoration: none; font-weight: 500;">➡️ 🗄️ Histórico</a>', unsafe_allow_html=True)

with col_usr:
    try:
        st.page_link(usr_path, label="👥 Usuários", icon="➡️")
    except:
        nome_pagina = usr_path.replace("pages/", "").replace(".py", "")
        st.markdown(f'<a href="{nome_pagina}" target="_top" style="display: block; text-align: center; background-color: #f0f2f6; border: 1px solid #d0d4dc; color: #31333F !important; padding: 0.5rem; border-radius: 0.5rem; text-decoration: none; font-weight: 500;">➡️ 👥 Usuários</a>', unsafe_allow_html=True)

with col_logout:
    if st.button("Sair", type="secondary", use_container_width=True):
        st.session_state.autenticado = False
        st.rerun()

st.markdown("---")

provedor_escolhido = st.selectbox("🧠 Motor de IA (Hub Multi-IA)", list(PROVEDORES_IA.keys()), key="provedor_ia_select")
cfg_provedor = PROVEDORES_IA[provedor_escolhido]
api_key = obter_chave_provedor(provedor_escolhido)
if not api_key:
    api_key = st.text_input(f"Chave da API ({cfg_provedor['secret']} não encontrada nos secrets)", type="password", placeholder=f"Cole sua chave de {provedor_escolhido} aqui...")
st.caption(f"Modelos utilizados (do mais capaz ao mais estável): {' → '.join(cfg_provedor['modelos'])}")

st.markdown("### 📥 Upload de Arquivos Normativos")
st.caption("Aceita Leis, Decretos, Resoluções, Enunciados, Portarias e demais atos normativos, em PDF.")
arquivos_enviados = st.file_uploader("Arraste todos os documentos (PDF) — um ato original e todos os seus derivativos", type=["pdf"], accept_multiple_files=True, key="uploader_lote")

# =====================================================================
# MOTOR DE PARSER HTML/XML (ÁRVORE SINTÁTICA AST)
# =====================================================================

class QuillParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.paragraphs = []
        self.current_html = []
        self.stack = []

    def _close_all_tags(self):
        res = ""
        for tags in reversed(self.stack):
            for t in reversed(tags):
                tag_name = t.split()[0]
                res += f"</{tag_name}>"
        return res

    def _open_all_tags(self):
        res = ""
        for tags in self.stack:
            for t in tags:
                res += f"<{t}>"
        return res

    def _break_paragraph(self):
        if self.current_html:
            p_text = "".join(self.current_html) + self._close_all_tags()
            if re.sub(r'<[^>]+>', '', p_text).strip():
                self.paragraphs.append(p_text.strip())
        self.current_html = []
        if self.stack:
            self.current_html.append(self._open_all_tags())

    def handle_starttag(self, tag, attrs):
        if tag in ('p', 'br', 'div'):
            self._break_paragraph()
            return

        attrs_dict = dict(attrs)
        style = attrs_dict.get('style', '').lower().replace(' ', '')
        cls = attrs_dict.get('class', '').lower()
        cor = attrs_dict.get('color', '').lower()
        
        added_tags = []
        if tag in ('b', 'strong') or 'font-weight:bold' in style or 'font-weight:700' in style:
            added_tags.append("b")
        if tag in ('i', 'em') or 'font-style:italic' in style:
            added_tags.append("i")
        if tag in ('s', 'strike', 'del') or 'text-decoration:line-through' in style or 'ql-strike' in cls:
            added_tags.append("strike")
        if ('color:rgb(230' in style or 'color:red' in style or 'color:#e6' in style or 'color:#f00' in style or 'color:#ff0000' in style) or (tag == 'font' and attrs_dict.get('color') in ('red', '#f00', '#ff0000')):
            added_tags.append('font color="red"')
        
        if added_tags:
            for t in added_tags:
                self.current_html.append(f"<{t}>")
            self.stack.append(added_tags)
        else:
            self.stack.append([])

    def handle_endtag(self, tag):
        if tag in ('p', 'br', 'div'):
            return 
        
        if self.stack:
            tags_to_close = self.stack.pop()
            for t in reversed(tags_to_close):
                tag_name = t.split()[0]
                self.current_html.append(f"</{tag_name}>")

    def handle_startendtag(self, tag, attrs):
        if tag in ('br',):
            self._break_paragraph()

    def handle_data(self, data):
        if not data: return
        data = data.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\xa0', '&nbsp;')
        self.current_html.append(data)

    def get_paragraphs(self):
        self._break_paragraph()
        return self.paragraphs

def ia_para_editor(texto):
    if not texto: return ""
    texto = texto.replace("<br/>", "</p><p>").replace("<br>", "</p><p>")
    if not texto.startswith("<p>"): texto = f"<p>{texto}</p>"
    texto = re.sub(r'<(strike|del)\b[^>]*>', '<s>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</(strike|del)>', '</s>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<font[^>]*color=[\'"]?(red|#f00|#ff0000|rgb\([^)]+\))[\'"]?[^>]*>', '<span style="color: rgb(230, 0, 0);">', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</font>', '</span>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<b\b[^>]*>', '<strong>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</b>', '</strong>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<i\b[^>]*>', '<em>', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</i>', '</em>', texto, flags=re.IGNORECASE)
    return texto.replace("<p></p>", "")

def editor_para_pdf(texto):
    if not texto: return ""
    parser = QuillParser()
    try:
        parser.feed(texto)
        return "<br/>".join(parser.get_paragraphs())
    except Exception:
        texto_limpo = re.sub(r'</?(span|div|p|ul|li|ol)[^>]*>', '', texto, flags=re.IGNORECASE)
        return texto_limpo

# =====================================================================
# LÓGICA DE NEGÓCIO E INTELIGÊNCIA ARTIFICIAL
# =====================================================================

SYSTEM_INSTRUCTION_LEGISTECNICA = """
Você é um Especialista Sênior em Técnica Legislativa do Poder Público brasileiro, apto a trabalhar com
QUALQUER espécie normativa: Leis, Decretos, Resoluções, Portarias, Enunciados, Instruções Normativas etc.
Nunca assuma que o documento é necessariamente uma Portaria. Regras obrigatórias:

1. FIDELIDADE ABSOLUTA: transcreva com exatidão o conteúdo de cada dispositivo, preservando formatação (<b>, <i>, quebras <br/>).
2. SEPARAÇÃO ESTRUTURAL OBRIGATÓRIA:
   - 'ementa': Resumo descritivo do objeto da norma.
   - 'preambulo': Autoridade expedidora e os Considerandos.
3. TABELAS: quando o dispositivo contiver uma tabela (identificada por marcadores [TABELA]...[/TABELA] no
   texto de origem), transcreva TODAS as linhas e colunas com fidelidade absoluta em 'tabela_alterada' e
   'tabela_consolidada' (uma lista de listas, uma sublista por linha, mantendo a ordem exata de colunas).
   NUNCA descreva a tabela em prosa, NUNCA a omita, e NUNCA resuma seu conteúdo — reproduza célula a célula,
   mesmo que a tabela seja grande. Se um ato alterador modifica um o conteúdo e dentro desse conteúdo existe uma tabela (acrescenta,
   remove ou muda linhas/colunas) ela deve ser taxada, com  <strike><font color="red"> Célula </font></strike>, 'tabela_alterada' deve conter a tabela NOVA e COMPLETA (com todas as
   linhas, alteradas ou não), e o campo 'texto_pos_tabela_alterada' deve trazer a nota
   "(Nova redação dada pelo Art. <N> da <TIPO> Nº <NÚMERO>/<SIGLA>, <DATA>)" logo abaixo da tabela.
   'tabela_consolidada' sempre reflete a versão vigente (mais recente) da tabela, a nova redação no caso de alterada deve vir após a tabela ou texto, verifique o que vem por último e após isos que deve vir a nova redação.
4. CRITÉRIO RIGOROSO DE ALTERAÇÃO E REVOGAÇÃO — formato EXATO e obrigatório (siga rigorosamente a
   pontuação, os parênteses e a ordem abaixo; NUNCA misture os dois casos):

   a) DISPOSITIVO ALTERADO (nova redação) — na versão ALTERADA (`texto_principal_alterada`), escreva em
      DUAS LINHAS separadas por quebra de parágrafo dupla `<br/><br/>` (nunca concatenadas na mesma linha):

      Linha 1 — identificador + texto ANTIGO INTEGRAL riscado em vermelho, seguido IMEDIATAMENTE (fora do
      risco, na mesma linha) da nota "(Alterada pelo Art. <N> da <TIPO> Nº <NÚMERO>/<SIGLA>, <DATA>)":
        <strike><font color="red">X - texto antigo integral, incluindo tabelas ...</font></strike> (Alterada pelo Art. 8 da
        PORTARIA Nº 1/PGJCG, de 01 JUNHO DE 2026) ou o texto continua.
      <br/><br/>
      Linha 2 — repita o MESMO identificador + a NOVA redação vigente por extenso, sem riscar, seguida da
      nota "(Redação dada pelo Art. <N> da <TIPO> Nº <NÚMERO>/<SIGLA>, <DATA>).":
        X - texto novo integral... (Redação dada pelo Art. 8 da PORTARIA Nº 1/PGJCG, de 01 JUNHO DE 2026).

      Exemplo completo (copie exatamente este padrão, adaptando o conteúdo):
        <strike><font color="red">X - apresentar ao final do período de instrutoria "Relatório das
        Atividades desenvolvidas durante o processo de Instrutoria", conforme modelo padrão.</font></strike>
        (Alterada pelo Art. 8 da PORTARIA Nº 1/PGJCG, de 01 JUNHO DE 2026)<br/><br/>X - apresentar trimestralmente
        relatórios de atividade ao(à) instruendo(a), conforme modelo anexo; (Redação dada pelo Art. 8 da
        PORTARIA Nº 1/PGJCG, de 01 JUNHO DE 2026).

      Na versão CONSOLIDADA (`texto_principal_consolidada`) do MESMO dispositivo, mostre SOMENTE a Linha 2
      (a nova redação), nunca a antiga, nunca riscada:
        X - apresentar trimestralmente relatórios de atividade ao(à) instruendo(a), conforme modelo anexo;
        (Redação dada pelo Art. 8 daPORTARIA Nº 1/PGJCG, de 01 JUNHO DE 2026).

   b) DISPOSITIVO REVOGADO — na versão ALTERADA, UMA ÚNICA LINHA (NÃO repita/acrescente uma segunda linha):
      identificador + texto INTEGRAL riscado em vermelho, seguido IMEDIATAMENTE da nota
      "(Revogado pelo Art. <N> da <TIPO> Nº <NÚMERO>/<SIGLA>, <DATA>);":
        <strike><font color="red">X - apresentar ao final do período de instrutoria "Relatório das
        Atividades desenvolvidas durante o processo de Instrutoria", conforme modelo padrão.</font></strike>
        (Revogado pelo Art. 8 da PORTARIA Nº 1/PGJCG, de 01 JUNHO DE 2026);

      Na versão CONSOLIDADA do mesmo dispositivo, mostre APENAS o identificador + a nota de revogação,
      sem repetir o texto revogado:
        X - (Revogado pelo Art. 8 da PORTARIA Nº 1/PGJCG, de 01 JUNHO DE 2026).

   c) As notas "(Alterada pelo ...)", "(Redação dada pelo ...)" e "(Revogado pelo ...)" vão SEMPRE
      embutidas diretamente no texto de 'texto_principal_alterada'/'texto_principal_consolidada' (não em
      campo separado), citando o ARTIGO ESPECÍFICO do ato alterador que promoveu a mudança — nunca cite
      só o ato inteiro sem o artigo. Preencha também 'nota_remissiva' com o mesmo trecho da citação (sem
      parênteses), só para fins de indexação/auditoria — mas isso é redundante ao texto, não substitui.
   d) NUNCA deixe de taxar o dispositivo correto, e NUNCA junte texto antigo e novo na mesma linha sem a
      quebra de parágrafo dupla `<br/><br/>` entre eles, exceto no caso de revogação (que é uma única linha).
5. GENERALIDADE: as regras acima valem para qualquer espécie normativa (Lei, Decreto, Resolução, Portaria,
   Enunciado, Instrução Normativa etc.) e para qualquer tipo de dispositivo (Artigo, Parágrafo, Parágrafo
   Único, Inciso, Alínea, Item).
"""

# ----------------- SCHEMA JSON PARA PROVEDORES SEM STRUCTURED OUTPUT NATIVO -----------------
def _prompt_schema_json(response_schema):
    esquema = response_schema.model_json_schema()
    return (
        "\n\nRESPONDA EXCLUSIVAMENTE COM UM OBJETO JSON VÁLIDO (sem markdown, sem ```json, sem comentários, "
        "sem texto antes ou depois) que obedeça RIGOROSAMENTE a este JSON Schema:\n"
        + json.dumps(esquema, ensure_ascii=False)
    )

def _extrair_json_bruto(texto):
    if not texto: raise Exception("Resposta vazia da IA.")
    t = texto.strip()
    t = re.sub(r'^```(json)?', '', t.strip(), flags=re.IGNORECASE).strip()
    t = re.sub(r'```$', '', t.strip()).strip()
    inicio = t.find('{')
    fim = t.rfind('}')
    if inicio == -1 or fim == -1: raise Exception("A IA não retornou um JSON reconhecível.")
    return t[inicio:fim + 1]

def _itens_para_texto_e_imagens(itens):
    """Separa a lista universal de conteúdo (strings e dicts {'tipo':'imagem',...})
    em texto concatenado + lista de imagens (mime, bytes), para provedores que não
    usam o formato de 'Part' do Gemini."""
    textos, imagens = [], []
    for it in itens:
        if isinstance(it, dict) and it.get("tipo") == "imagem":
            imagens.append((it["mime"], it["dados"]))
        elif isinstance(it, str):
            textos.append(it)
    return "\n\n".join(textos), imagens

def _itens_para_parts_gemini(itens):
    partes = []
    for it in itens:
        if isinstance(it, dict) and it.get("tipo") == "imagem":
            partes.append(types.Part.from_bytes(data=it["dados"], mime_type=it["mime"]))
        elif isinstance(it, str):
            partes.append(it)
    return partes

def _chamar_gemini(chave, itens, response_schema, thinking_level, modelos):
    client = genai.Client(api_key=chave)
    config = types.GenerateContentConfig(
        response_mime_type="application/json",
        response_schema=response_schema,
        system_instruction=SYSTEM_INSTRUCTION_LEGISTECNICA,
        thinking_config=types.ThinkingConfig(thinking_level=thinking_level),
    )
    contents = _itens_para_parts_gemini(itens)
    ultimo_erro = None
    for modelo in modelos:
        cota_diaria_esgotada = False
        for tentativa in range(1, 4):
            try:
                resp = client.models.generate_content(model=modelo, contents=contents, config=config)
                _validar_resposta_gemini(resp)
                dados = json.loads(resp.text)
                return response_schema.model_validate(dados)
            except Exception as e:
                ultimo_erro = e
                erro_str = str(e).upper()
                if "PERDAY" in erro_str.replace(" ", "") or "FREE_TIER" in erro_str or "GENERATEREQUESTSPERDAY" in erro_str.replace(" ", ""):
                    # Cota DIÁRIA (não por minuto) esgotada — retry não adianta, o limite só
                    # reseta no dia seguinte. Pula direto para o próximo modelo/provedor.
                    st.toast(f"⚠️ Cota diária do {modelo} esgotada (free tier). Pulando para o próximo modelo...", icon="📅")
                    cota_diaria_esgotada = True
                    break
                elif "429" in erro_str or "RESOURCE_EXHAUSTED" in erro_str or "503" in erro_str or "UNAVAILABLE" in erro_str:
                    if tentativa < 3:
                        tempo_espera = min(tentativa * 3, 10)
                        st.toast(f"⚡ Fila no Google ({modelo}). Tentativa {tentativa}/3. Aguardando {tempo_espera}s...", icon="⏳")
                        time.sleep(tempo_espera)
                        continue
                    st.toast(f"⚡ Tempo esgotado no {modelo}. Mudando para o próximo...", icon="🔄")
                    break
                elif "404" in erro_str or "NOT_FOUND" in erro_str or "400" in erro_str:
                    st.toast(f"⚠️ Modelo {modelo} indisponível. Pulando...", icon="⏭️")
                    break
                else:
                    raise e
        if cota_diaria_esgotada:
            continue
    raise Exception(f"Google Gemini: todos os modelos falharam (cota diária pode estar esgotada no free tier). Último erro: {ultimo_erro}")

def _validar_resposta_gemini(resp):
    candidatos = getattr(resp, "candidates", None) or []
    if candidatos:
        finish = getattr(candidatos[0], "finish_reason", None)
        finish_str = str(finish) if finish else ""
        if "MAX_TOKENS" in finish_str: raise Exception("A resposta da IA foi cortada por limite de tokens.")
        if "SAFETY" in finish_str or "PROHIBITED" in finish_str: raise Exception("Bloqueado por política de segurança.")
    if not getattr(resp, "text", None): raise Exception("Resposta vazia da IA.")

def _montar_mensagens_openai_like(itens, response_schema):
    texto, imagens = _itens_para_texto_e_imagens(itens)
    texto += _prompt_schema_json(response_schema)
    conteudo_usuario = [{"type": "text", "text": texto}]
    for mime, dados in imagens:
        b64 = base64.b64encode(dados).decode()
        conteudo_usuario.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}})
    mensagens = [
        {"role": "system", "content": SYSTEM_INSTRUCTION_LEGISTECNICA},
        {"role": "user", "content": conteudo_usuario if imagens else texto},
    ]
    return mensagens

def _chamar_groq(chave, itens, response_schema, modelos):
    if Groq is None: raise Exception("Biblioteca 'groq' não instalada no servidor.")
    client = Groq(api_key=chave)
    mensagens = _montar_mensagens_openai_like(itens, response_schema)
    ultimo_erro = None
    for modelo in modelos:
        for tentativa in range(1, 4):
            try:
                resp = client.chat.completions.create(
                    model=modelo, messages=mensagens,
                    response_format={"type": "json_object"}, temperature=0.2,
                )
                bruto = _extrair_json_bruto(resp.choices[0].message.content)
                return response_schema.model_validate(json.loads(bruto))
            except Exception as e:
                ultimo_erro = e
                erro_str = str(e).upper()
                if "429" in erro_str or "RATE_LIMIT" in erro_str or "503" in erro_str:
                    if tentativa < 3:
                        tempo_espera = min(tentativa * 3, 10)
                        st.toast(f"⚡ Fila na Groq ({modelo}). Tentativa {tentativa}/3. Aguardando {tempo_espera}s...", icon="⏳")
                        time.sleep(tempo_espera)
                        continue
                    break
                elif "404" in erro_str or "NOT_FOUND" in erro_str or isinstance(e, (ValidationError, json.JSONDecodeError)) or "JSON" in erro_str.upper() or "não retornou" in str(e):
                    st.toast(f"⚠️ {modelo} indisponível/formato inválido. Pulando...", icon="⏭️")
                    break
                else:
                    raise e
    raise Exception(f"Groq: todos os modelos falharam. Último erro: {ultimo_erro}")

def _chamar_openrouter(chave, itens, response_schema, modelos):
    if OpenAI is None: raise Exception("Biblioteca 'openai' não instalada no servidor.")
    client = OpenAI(api_key=chave, base_url="https://openrouter.ai/api/v1")
    mensagens = _montar_mensagens_openai_like(itens, response_schema)
    ultimo_erro = None
    for modelo in modelos:
        for tentativa in range(1, 4):
            try:
                resp = client.chat.completions.create(
                    model=modelo, messages=mensagens,
                    response_format={"type": "json_object"}, temperature=0.2,
                )
                bruto = _extrair_json_bruto(resp.choices[0].message.content)
                return response_schema.model_validate(json.loads(bruto))
            except Exception as e:
                ultimo_erro = e
                erro_str = str(e).upper()
                if "429" in erro_str or "RATE_LIMIT" in erro_str or "503" in erro_str:
                    if tentativa < 3:
                        tempo_espera = min(tentativa * 3, 10)
                        st.toast(f"⚡ Fila no OpenRouter ({modelo}). Tentativa {tentativa}/3. Aguardando {tempo_espera}s...", icon="⏳")
                        time.sleep(tempo_espera)
                        continue
                    break
                elif "404" in erro_str or "NOT_FOUND" in erro_str or isinstance(e, (ValidationError, json.JSONDecodeError)) or "JSON" in erro_str.upper() or "não retornou" in str(e):
                    st.toast(f"⚠️ {modelo} indisponível/formato inválido. Pulando...", icon="⏭️")
                    break
                else:
                    raise e
    raise Exception(f"OpenRouter: todos os modelos falharam. Último erro: {ultimo_erro}")

def _chamar_mistral(chave, itens, response_schema, modelos):
    if Mistral is None: raise Exception("Biblioteca 'mistralai' não instalada no servidor.")
    client = Mistral(api_key=chave)
    mensagens = _montar_mensagens_openai_like(itens, response_schema)
    ultimo_erro = None
    for modelo in modelos:
        for tentativa in range(1, 4):
            try:
                resp = client.chat.complete(
                    model=modelo, messages=mensagens,
                    response_format={"type": "json_object"}, temperature=0.2,
                )
                bruto = _extrair_json_bruto(resp.choices[0].message.content)
                return response_schema.model_validate(json.loads(bruto))
            except Exception as e:
                ultimo_erro = e
                erro_str = str(e).upper()
                if "429" in erro_str or "CAPACITY" in erro_str or "503" in erro_str:
                    if tentativa < 3:
                        tempo_espera = min(tentativa * 3, 10)
                        st.toast(f"⚡ Fila na Mistral ({modelo}). Tentativa {tentativa}/3. Aguardando {tempo_espera}s...", icon="⏳")
                        time.sleep(tempo_espera)
                        continue
                    break
                elif "404" in erro_str or "NOT_FOUND" in erro_str or isinstance(e, (ValidationError, json.JSONDecodeError)) or "JSON" in erro_str.upper() or "não retornou" in str(e):
                    st.toast(f"⚠️ {modelo} indisponível/formato inválido. Pulando...", icon="⏭️")
                    break
                else:
                    raise e
    raise Exception(f"Mistral AI: todos os modelos falharam. Último erro: {ultimo_erro}")

def _chamar_por_motor(motor, chave, itens, response_schema, thinking_level, modelos):
    if motor == "gemini":
        return _chamar_gemini(chave, itens, response_schema, thinking_level, modelos)
    elif motor == "groq":
        return _chamar_groq(chave, itens, response_schema, modelos)
    elif motor == "openrouter":
        return _chamar_openrouter(chave, itens, response_schema, modelos)
    elif motor == "mistral":
        return _chamar_mistral(chave, itens, response_schema, modelos)
    raise Exception(f"Provedor desconhecido: {motor}")

def executar_com_fallback(chave, itens, response_schema, provedor, thinking_level="high"):
    """Despacha a chamada para o provedor de IA escolhido (Gemini, Groq, OpenRouter
    ou Mistral), sempre validando o resultado contra o schema Pydantic esperado.
    Se o provedor escolhido falhar por completo (ex.: cota diária esgotada no free
    tier), tenta automaticamente os outros provedores que tiverem chave configurada
    nos secrets, na ordem do Hub Multi-IA, antes de desistir."""
    cfg = PROVEDORES_IA[provedor]
    try:
        resultado = _chamar_por_motor(cfg["motor"], chave, itens, response_schema, thinking_level, cfg["modelos"])
    except Exception as erro_provedor_escolhido:
        outros = [p for p in PROVEDORES_IA if p != provedor]
        ultimo_erro = erro_provedor_escolhido
        resultado = None
        for nome_alt in outros:
            chave_alt = obter_chave_provedor(nome_alt)
            if not chave_alt:
                continue
            try:
                st.toast(f"🔀 {provedor} indisponível. Tentando automaticamente com {nome_alt}...", icon="🔁")
                cfg_alt = PROVEDORES_IA[nome_alt]
                resultado = _chamar_por_motor(cfg_alt["motor"], chave_alt, itens, response_schema, thinking_level, cfg_alt["modelos"])
                break
            except Exception as e2:
                ultimo_erro = e2
                continue
        if resultado is None:
            raise Exception(f"{provedor} falhou e nenhum provedor alternativo configurado deu certo. Último erro: {ultimo_erro}")

    class _RespCompat:
        def __init__(self, obj): self.text = obj.model_dump_json()
    return _RespCompat(resultado)

def converter_para_iso(data_str):
    if not data_str: return None
    data_str = data_str.strip()
    if re.match(r'^\d{4}-\d{2}-\d{2}$', data_str): return data_str
    match_br = re.match(r'^(\d{2})/(\d{2})/(\d{4})$', data_str)
    if match_br:
        d, m, a = match_br.groups()
        return f"{a}-{m}-{d}"
    try: return datetime.strptime(data_str, "%d/%m/%Y").strftime("%Y-%m-%d")
    except: return None

def extrair_conteudo_multimodal(file_bytes, nome_arquivo):
    """Retorna uma lista de itens universais (str = texto, ou dict {'tipo':'imagem',
    'mime':..., 'dados':...}) — formato independente de provedor de IA, convertido
    depois para o formato específico de cada motor (Gemini/Groq/OpenRouter/Mistral).
    Detecta tabelas com fitz.Table Finder e as marca explicitamente com
    [TABELA]...[/TABELA] para a IA nunca perder a estrutura tabular."""
    if nome_arquivo.lower().endswith(".docx"): return [f"ARQUIVO DOCX: {nome_arquivo}"]
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        html_text = f"CONTEÚDO DO ARQUIVO {nome_arquivo}:\n\n"
        caracteres_uteis = 0
        for page_num, page in enumerate(doc):
            html_text += f"=== PÁGINA {page_num + 1} ===\n"

            tabelas_bbox = []
            try:
                tab_finder = page.find_tables()
                for tabela in tab_finder.tables:
                    linhas = tabela.extract()
                    if not linhas: continue
                    caracteres_uteis += sum(len(str(c or "")) for linha in linhas for c in linha)
                    tabelas_bbox.append(fitz.Rect(tabela.bbox))
                    html_text += "[TABELA]\n"
                    for linha in linhas:
                        html_text += " | ".join((str(c).strip() if c is not None else "") for c in linha) + "\n"
                    html_text += "[/TABELA]\n<br/>\n"
            except Exception:
                pass

            blocks = page.get_text("dict", sort=True).get("blocks", [])
            for b in blocks:
                if b.get('type') != 0: continue
                bloco_rect = fitz.Rect(b.get("bbox", (0, 0, 0, 0)))
                if any(bloco_rect.intersects(tb) for tb in tabelas_bbox):
                    continue  # já capturado como [TABELA] acima — evita duplicar/embaralhar
                bloco_linhas = ""
                for l in b.get("lines", []):
                    linha_span = ""
                    for s in l.get("spans", []):
                        texto = s.get("text", "")
                        if not texto: continue
                        caracteres_uteis += len(texto.strip())
                        flags = s.get("flags", 0)
                        if flags & 2**4: texto = f"<b>{texto}</b>"
                        if flags & 2**1: texto = f"<i>{texto}</i>"
                        linha_span += texto
                    if linha_span.strip(): bloco_linhas += linha_span + " "
                if bloco_linhas.strip(): html_text += bloco_linhas.strip() + "<br/>\n"
            html_text += "<br/>\n"

        if caracteres_uteis < 30 * max(doc.page_count, 1):
            partes = [f"ARQUIVO {nome_arquivo} É UM DOCUMENTO ESCANEADO. Leia o conteúdo visualmente, inclusive tabelas:"]
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                partes.append({"tipo": "imagem", "mime": "image/jpeg", "dados": pix.tobytes("jpg", jpg_quality=78)})
            return partes

        return [html_text]
    except Exception as e: return [f"Erro ao extrair PDF {nome_arquivo}: {str(e)}"]

# ----------------- ESTRUTURAS PYDANTIC -----------------
class ArquivoClassificado(BaseModel):
    nome_arquivo_upload: str
    tipo: str = Field(description="'Base' ou 'Alteradora'")
    grupo_id: int = Field(description="Identificador da família normativa (comece em 1).")
    nome_padronizado_identificado: str = Field(description="Nome padronizado da norma (tipo, número, órgão e data)")
    data_oficial_iso: str = Field(description="Data formatada estritamente em YYYY-MM-DD.")
    ato_base_referenciado_tipo: Optional[str] = Field(default=None, description="APENAS para 'Alteradora' cujo ato original NÃO está presente neste lote: o tipo do ato que ela declara alterar/revogar (ex.: 'PORTARIA'). Deixe vazio se o ato base está no próprio lote ou se tipo='Base'.")
    ato_base_referenciado_numero: Optional[str] = Field(default=None, description="APENAS para 'Alteradora' cujo ato original NÃO está presente neste lote: o número/identificador do ato que ela declara alterar/revogar (ex.: '158/PGJM'), extraído do próprio texto da alteradora (normalmente citado no preâmbulo/ementa). Deixe vazio se o ato base está no próprio lote ou se tipo='Base'.")

class TriagemDocumentos(BaseModel): arquivos: List[ArquivoClassificado]

class MetadadosNorma(BaseModel):
    tipo_documento: str; numero_documento: str; orgao_emissor: str; data_assinatura: str; nome_padronizado: str

class Dispositivo(BaseModel):
    tipo: str; texto_principal_alterada: str; texto_principal_consolidada: str; is_tabela: bool
    tabela_alterada: Optional[List[List[str]]] = None; tabela_consolidada: Optional[List[List[str]]] = None
    texto_pos_tabela_alterada: Optional[str] = None; texto_pos_tabela_consolidada: Optional[str] = None
    nota_remissiva: Optional[str] = Field(default="", description="Apenas o trecho da citação, sem prefixo e sem parênteses, ex.: 'Art. 8 da PORTARIA Nº 1/2026 - PGJ/CG'. A citação completa com o prefixo ('Alterada pelo', 'Redação dada pelo' ou 'Revogado pelo') já vai embutida em texto_principal_alterada/consolidada — este campo é só para indexação.")

class Consolidacao(BaseModel):
    arquivos_originais_identificados: List[str]; arquivos_alteradores_identificados: List[str]
    norma_base: MetadadosNorma; normas_alteradoras: List[MetadadosNorma]
    cabecalho_complemento: str; orgaos_emissores: str; titulo_portaria: str; ementa: str; preambulo: str
    assinatura_nome: str; assinatura_cargo: str; dispositivos: List[Dispositivo]

class AnaliseGlobal(BaseModel):
    consolidacoes_geradas: List[Consolidacao]; arquivos_nao_alterados: List[str]

def limpar_texto_ia(texto):
    if not texto: return ""
    return re.sub(r' {2,}', ' ', str(texto)).strip()

def injetar_nota_remissiva(texto, nota):
    if nota and nota.strip():
        n_sem_parenteses = nota.strip("()").strip()
        n_fmt = f"({n_sem_parenteses})"
        
        texto_puro = re.sub(r'<[^>]+>', '', texto if texto else '')
        if n_sem_parenteses.lower() in texto_puro.lower(): return texto 
        
        if texto:
            texto_limpo = re.sub(r'(<br/?>|\s)+$', '', texto).strip()
            return f'{texto_limpo} &nbsp;<span style="color: red;">{n_fmt}</span>'
        return f'<span style="color: red;">{n_fmt}</span>'
    return texto

def resgatar_memoria():
    memoria = ""
    if supabase:
        try:
            res = supabase.table("memoria_de_correcoes").select("*").order("id", desc=True).limit(5).execute()
            if res.data:
                memoria = "\n\n⚠️ HISTÓRICO DE CORREÇÕES (Não repita os erros da IA):\n"
                for m in res.data: memoria += f"- Erro: {m['texto_ia']}\n- Correção: {m['texto_corrigido']}\n\n"
        except: pass
    return memoria

# =====================================================================
# EDITOR ÚNICO POR TAGS — serializa/parseia o documento inteiro (versão
# ALTERADA) como um único texto marcado, e deriva a versão CONSOLIDADA
# programaticamente a partir dele (nunca editada separadamente).
# =====================================================================

def _localizar_base_no_banco(tipo_ref, numero_ref):
    """Quando uma alteradora chega isolada (sem o ato original no mesmo lote),
    procura no Supabase um ato base já cadastrado que corresponda ao que ela
    declara alterar/revogar."""
    if not supabase or not numero_ref or not str(numero_ref).strip():
        return None
    try:
        numero_limpo = str(numero_ref).strip()
        query = supabase.table("portarias_base").select("id, nome_padronizado, tipo_documento, numero_documento, documento_consolidado_json")
        res = query.ilike("numero_documento", f"%{numero_limpo}%").execute()
        candidatos = res.data or []
        if not candidatos and tipo_ref:
            res2 = query.ilike("nome_padronizado", f"%{numero_limpo}%").execute()
            candidatos = res2.data or []
        if not candidatos:
            return None
        if tipo_ref:
            for c in candidatos:
                if str(c.get('tipo_documento', '')).strip().lower() == str(tipo_ref).strip().lower():
                    return c
        return candidatos[0]
    except Exception:
        return None

def analisar_lote_arquivos(arquivos, key, provedor):
    memoria_aprendida = resgatar_memoria()

    textos_extraidos = {}
    with ThreadPoolExecutor(max_workers=min(8, max(1, len(arquivos)))) as ex:
        futuros = {submit_com_contexto(ex, extrair_conteudo_multimodal, arq.getvalue(), arq.name): arq.name for arq in arquivos}
        for fut in as_completed(futuros):
            textos_extraidos[futuros[fut]] = fut.result()

    contents_triagem = [f"Analise os documentos. Agrupe cada ato original com seus derivativos presentes neste lote. Se uma Alteradora citar um ato original que NÃO está entre os arquivos deste lote, preencha ato_base_referenciado_tipo/numero com o que ela declara alterar/revogar, para localização posterior no banco de dados. ARQUIVOS: {', '.join(textos_extraidos.keys())}"]
    for partes in textos_extraidos.values(): contents_triagem.extend(partes)
    resp_triagem = executar_com_fallback(key, contents_triagem, TriagemDocumentos, provedor, thinking_level="low")
    triagem_dados = json.loads(resp_triagem.text).get("arquivos", [])

    grupos = {}
    for a in triagem_dados: grupos.setdefault(a.get('grupo_id', 0), []).append(a)

    grupos_validos = []
    consolidacoes_geradas, arquivos_nao_alterados = [], []
    for grupo_id, itens in grupos.items():
        arquivo_base = next((a for a in itens if a['tipo'] == 'Base'), None)
        arquivos_alteradores = sorted([a for a in itens if a['tipo'] == 'Alteradora'], key=lambda x: x['data_oficial_iso'])

        if not arquivo_base and not arquivos_alteradores: continue
        if not arquivo_base:
            base_reconstruida = None
            for alt in arquivos_alteradores:
                candidato = _localizar_base_no_banco(alt.get('ato_base_referenciado_tipo'), alt.get('ato_base_referenciado_numero'))
                if candidato:
                    base_reconstruida = candidato
                    break
            if base_reconstruida:
                arquivo_base = {
                    "nome_arquivo_upload": None,
                    "tipo": "Base",
                    "nome_padronizado_identificado": base_reconstruida.get("nome_padronizado", ""),
                    "data_oficial_iso": "",
                    "_reconstruida_do_banco": True,
                }
                grupos_validos.append((arquivo_base, arquivos_alteradores))
            else:
                arquivos_nao_alterados.extend([a['nome_arquivo_upload'] for a in arquivos_alteradores])
            continue
        grupos_validos.append((arquivo_base, arquivos_alteradores))

    if grupos_validos:
        with ThreadPoolExecutor(max_workers=min(4, len(grupos_validos))) as ex:
            futuros = {}
            for arquivo_base, arquivos_alteradores in grupos_validos:
                st.toast(f"⚙️ Processando: {arquivo_base.get('nome_padronizado_identificado')}...", icon="⏳")
                fut = submit_com_contexto(ex, _processar_cascata_grupo, key, provedor, arquivo_base, arquivos_alteradores, textos_extraidos, memoria_aprendida)
                futuros[fut] = (arquivo_base, arquivos_alteradores)
            for fut in as_completed(futuros):
                arquivo_base, arquivos_alteradores = futuros[fut]
                try:
                    resultado, mensagens = fut.result()
                    consolidacoes_geradas.append(resultado)
                    for tipo_msg, texto_msg in mensagens:
                        if tipo_msg == "info": st.info(texto_msg)
                        elif tipo_msg == "warning": st.warning(texto_msg)
                except Exception as e:
                    st.error(f"❌ Falha em '{arquivo_base.get('nome_padronizado_identificado')}': {e}")
                    if arquivo_base.get('nome_arquivo_upload'):
                        arquivos_nao_alterados.append(arquivo_base['nome_arquivo_upload'])
                    arquivos_nao_alterados.extend([a['nome_arquivo_upload'] for a in arquivos_alteradores])

    return {"consolidacoes_geradas": consolidacoes_geradas, "arquivos_nao_alterados": arquivos_nao_alterados}

def _consultar_estado_e_historico(nome_padrao):
    """Verifica no Supabase se este ato original já tem processamento salvo e
    quais alteradoras/revogadoras já foram aplicadas a ele anteriormente."""
    if not supabase or not nome_padrao:
        return None, []
    try:
        res_bd = supabase.table("portarias_base").select("id, documento_consolidado_json").eq("nome_padronizado", nome_padrao).execute()
        if not res_bd.data:
            return None, []
        base_id = res_bd.data[0]['id']
        estado = res_bd.data[0].get("documento_consolidado_json")
        res_alt = supabase.table("portarias_alteradoras").select("nome_padronizado").eq("portaria_base_id", base_id).execute()
        ja_processadas = [r['nome_padronizado'] for r in (res_alt.data or []) if r.get('nome_padronizado')]
        return (json.dumps(estado) if estado else None), ja_processadas
    except Exception:
        return None, []

def _processar_cascata_grupo(key, provedor, arquivo_base, arquivos_alteradores, textos_extraidos, memoria_aprendida):
    nome_padrao = arquivo_base.get('nome_padronizado_identificado', '')
    reconstruida = bool(arquivo_base.get('_reconstruida_do_banco'))
    estado_json_atual, ja_processadas = _consultar_estado_e_historico(nome_padrao)
    mensagens = []

    if reconstruida:
        detalhe = f" com {len(ja_processadas)} derivação(ões) já aplicada(s) ({', '.join(ja_processadas)})" if ja_processadas else ""
        mensagens.append(("info", f"📎 Os arquivos enviados alteram/revogam o ato '{nome_padrao}', já cadastrado no banco{detalhe}. Recomendamos anexar também o arquivo ORIGINAL de '{nome_padrao}' em um novo envio para garantir a máxima fidelidade; por ora, o processamento usará o estado já consolidado salvo no banco de dados."))
    elif estado_json_atual is not None:
        detalhe = f" ({', '.join(ja_processadas)})" if ja_processadas else ""
        mensagens.append(("info", f"🧠 '{nome_padrao}' já possui histórico no banco: {len(ja_processadas)} alteração(ões)/revogação(ões) processada(s) anteriormente{detalhe}."))

    if reconstruida and estado_json_atual is None:
        mensagens.append(("warning", f"⚠️ '{nome_padrao}' foi localizado no banco, mas sem conteúdo consolidado salvo. Envie também o arquivo ORIGINAL de '{nome_padrao}' junto com as alteradoras para que o processamento seja possível."))
        raise Exception(f"Ato base '{nome_padrao}' localizado no banco sem conteúdo salvo — reenvie junto com o arquivo original.")

    ja_processadas_lower = {j.lower() for j in ja_processadas}
    alteradoras_para_aplicar = []
    for alt in arquivos_alteradores:
        nome_alt = alt.get('nome_padronizado_identificado', '')
        if nome_alt and nome_alt.lower() in ja_processadas_lower:
            mensagens.append(("warning", f"⚠️ '{nome_alt}' já havia sido processada e aplicada anteriormente a '{nome_padrao}' — não será reaplicada agora para evitar duplicar a alteração/revogação."))
        else:
            alteradoras_para_aplicar.append(alt)
    alteradoras_para_aplicar.sort(key=lambda x: x.get('data_oficial_iso') or '')

    if not alteradoras_para_aplicar:
        if estado_json_atual:
            return json.loads(estado_json_atual), mensagens
        conteudo_loop = ["Texto Base:"] + textos_extraidos[arquivo_base['nome_arquivo_upload']]
        resp_loop = executar_com_fallback(key, conteudo_loop + ["Estruture o documento separando a ementa do preâmbulo e aplicando rigorosamente o mapeamento de dispositivos, incluindo tabelas quando houver." + memoria_aprendida], Consolidacao, provedor)
        return json.loads(resp_loop.text), mensagens

    resp_loop = None
    for i, alt in enumerate(alteradoras_para_aplicar):
        conteudo_loop = []
        if estado_json_atual:
            conteudo_loop.append(f"ESTADO ATUAL (JSON):\n{estado_json_atual}")
        elif i == 0:
            conteudo_loop.append("DOCUMENTO BASE ORIGINAL:")
            conteudo_loop.extend(textos_extraidos[arquivo_base['nome_arquivo_upload']])

        conteudo_loop.append(f"ATO ALTERADOR/REVOGADOR Nº {i+1} DE {len(alteradoras_para_aplicar)} A SER APLICADO, EM ORDEM CRONOLÓGICA DO MAIS ANTIGO PARA O MAIS NOVO ({alt['nome_arquivo_upload']}):")
        conteudo_loop.extend(textos_extraidos[alt['nome_arquivo_upload']])
        prompt_loop = f"""
        Aplique o ato alterador/revogador cruzando detalhadamente com o ato base (seja Lei, Decreto,
        Resolução, Enunciado, Portaria ou qualquer outra espécie normativa).
        Obrigatório: siga EXATAMENTE o formato de citação e a estrutura de parágrafos definidos nas regras
        do sistema (item 4) — "(Alterada pelo Art. N da TIPO Nº NÚMERO/ANO - SIGLA)" na linha riscada,
        "(Redação dada pelo Art. N da TIPO Nº NÚMERO/ANO - SIGLA)" na linha nova, ou "(Revogado pelo Art. N
        da TIPO Nº NÚMERO/ANO - SIGLA)" para revogação (sem repetir linha). Preserve <b>/<i> e tabelas
        ([TABELA]...[/TABELA]) com fidelidade absoluta, redesenhando a tabela inteira quando alterada.
        {memoria_aprendida}
        """
        conteudo_loop.append(prompt_loop)
        resp_loop = executar_com_fallback(key, conteudo_loop, Consolidacao, provedor)
        estado_json_atual = resp_loop.text
    return json.loads(resp_loop.text), mensagens

# =====================================================================
# EXPORTAÇÃO (HTML UNIVERSAL -> WEASYPRINT PDF -> DOCX AST)
# =====================================================================

def gerar_html_dinamico(consolidacao_dict, tipo_versao):
    comp = consolidacao_dict.get("cabecalho_complemento", "")
    titulo_doc = f"VERSÃO {'ALTERADA' if tipo_versao=='alterada' else 'CONSOLIDADA'} - {comp}"
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>{titulo_doc}</title>
        <style>
            @page {{ size: A4; margin: 2.5cm 2cm; }}
            /* FONTE EXATA EM 11pt / 11px */
            body {{ font-family: 'Times New Roman', Times, serif; font-size: 11pt; line-height: 1.5; text-align: justify; }}
            .topo {{ text-align: center; color: #444; font-size: 10pt; font-weight: bold; margin-bottom: 20px; text-transform: uppercase; }}
            .brasao {{ text-align: center; margin-bottom: 10px; }}
            .brasao img {{ width: 45px; height: auto; display: block; margin: 0 auto; }}
            .orgaos {{ text-align: center; font-weight: bold; margin-bottom: 25px; }}
            .titulo {{ text-align: center; font-weight: bold; margin-bottom: 20px; }}
            
            /* EMENTA: Alinhada à direita e recuada */
            .ementa {{ text-align: justify; margin-left: 45%; margin-bottom: 25px; font-weight: normal; }}
            
            /* PREÂMBULO E CONSIDERANDOS: Alinhados à esquerda, sem recuo de parágrafo */
            .preambulo {{ text-align: justify; margin-bottom: 12px; text-indent: 0; }}
            
            /* DISPOSITIVOS (Artigos): Recuo padrão de parágrafo legal (40px) */
            .dispositivo {{ text-align: justify; text-indent: 40px; margin-bottom: 12px; }}
            
            .capitulo {{ text-align: center; font-weight: bold; margin-top: 20px; margin-bottom: 12px; text-transform: uppercase; }}
            .assinatura {{ text-align: center; font-weight: bold; margin-top: 50px; margin-bottom: 20px; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 15px; }}
            td, th {{ border: 1px solid black; padding: 6px; text-align: left; vertical-align: middle; }}
            
            /* GARANTIA ABSOLUTA DE RENDERIZAÇÃO DO TAXADO E CORES */
            strike, s, del {{ text-decoration: line-through; }}
            b, strong {{ font-weight: bold; }}
            i, em {{ font-style: italic; }}
            font[color="red"], span[style*="color: red"], span[style*="color:rgb(230"] {{ color: red !important; }}
        </style>
    </head>
    <body>
        <div class="topo">{titulo_doc}</div>
    """
    
    if os.path.exists("brasao.png"):
        import base64
        with open("brasao.png", "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode()
            html += f"<div class='brasao'><img src='data:image/png;base64,{encoded_string}'/></div>"

    html += f"""
        <div class="orgaos">{limpar_texto_ia(consolidacao_dict.get("orgaos_emissores") or "").replace('<br/>', '<br>')}</div>
        <div class="titulo">{limpar_texto_ia(consolidacao_dict.get("titulo_portaria") or "").replace('<br/>', '<br>')}</div>
        <div class="ementa">{limpar_texto_ia(consolidacao_dict.get("ementa") or "").replace('<br/>', '<br>')}</div>
        <div class="preambulo">{limpar_texto_ia(consolidacao_dict.get("preambulo") or "").replace('<br/>', '<br>')}</div>
    """
    
    for item in consolidacao_dict.get("dispositivos", []):
        t = (item.get("tipo") or "").lower()
        t_prin = injetar_nota_remissiva(item.get(f"texto_principal_{tipo_versao}"), item.get("nota_remissiva") if not item.get("is_tabela") else "")
        
        if "capitulo" in t or "anexo" in t:
            html += f"<div class='capitulo'>{t_prin}</div>"
            if not item.get("is_tabela"):
                continue
        else:
            if t_prin:
                for p in t_prin.split("<br/>"):
                    if p.strip():
                        html += f"<div class='dispositivo'>{p.strip()}</div>"
            
        if item.get("is_tabela"):
            linhas = item.get(f"tabela_{tipo_versao}") or []
            if linhas:
                html += "<table>"
                for linha in linhas:
                    html += "<tr>"
                    for celula in linha: html += f"<td>{editor_para_pdf(celula)}</td>"
                    html += "</tr>"
                html += "</table>"
                
            t_pos = injetar_nota_remissiva(item.get(f"texto_pos_tabela_{tipo_versao}"), item.get("nota_remissiva"))
            if t_pos: 
                for p in t_pos.split("<br/>"):
                    if p.strip():
                        html += f"<div class='dispositivo'>{p.strip()}</div>"
                
    html += f"<div class='assinatura'>{limpar_texto_ia(consolidacao_dict.get('assinatura_nome') or '')}<br>{limpar_texto_ia(consolidacao_dict.get('assinatura_cargo') or '')}</div>"
    html += "</body></html>"
    return html

def gerar_pdf_dinamico(consolidacao_dict, tipo_versao):
    html_str = gerar_html_dinamico(consolidacao_dict, tipo_versao)
    if not HAS_WEASYPRINT:
        raise Exception(
            "WeasyPrint não está disponível neste servidor (faltam bibliotecas nativas do "
            "sistema — Pango/Cairo/GDK-Pixbuf). Adicione um arquivo packages.txt na raiz do "
            "repositório com as dependências do sistema e faça um novo deploy."
        )
    buffer = io.BytesIO()
    WeasyHTML(string=html_str).write_pdf(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def aplicar_html_no_docx(p, texto_html):
    texto_html = texto_html.replace("&nbsp;", "\xa0")
    tokens = re.split(r'(<[^>]+>)', texto_html)
    is_bold = is_strike = is_red = is_italic = False
    
    for token in tokens:
        if not token: continue
        t = token.lower()
        if t.startswith('<b') and not t.startswith('<br'): is_bold = True
        elif t == '</b>': is_bold = False
        elif t.startswith('<i'): is_italic = True
        elif t == '</i>': is_italic = False
        elif t.startswith('<strike') or t.startswith('<s') and not t.startswith('<span'): is_strike = True
        elif t == '</strike>' or t == '</s>': is_strike = False
        elif t.startswith('<font') and ('red' in t or '#f00' in t or '#e6' in t): is_red = True
        elif t.startswith('<span') and ('red' in t or '#f00' in t or '#e6' in t): is_red = True
        elif t == '</font>' or t == '</span>': is_red = False
        elif token.startswith('<'): pass
        else:
            token = unescape(token)
            run = p.add_run(token)
            run.font.name, run.font.size = 'Times New Roman', Pt(11)
            if is_bold: run.bold = True
            if is_italic: run.italic = True
            if is_strike: run.font.strike = True
            if is_red: run.font.color.rgb = RGBColor(230, 0, 0)

def gerar_docx_dinamico(consolidacao_dict, tipo_versao):
    doc = docx.Document()
    for section in doc.sections: section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

    ph = doc.add_paragraph(); ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = ph.add_run(f"VERSÃO {'ALTERADA' if tipo_versao=='alterada' else 'CONSOLIDADA'} - {consolidacao_dict.get('cabecalho_complemento', '')}")
    rh.font.name, rh.font.size, rh.bold, rh.font.color.rgb = 'Times New Roman', Pt(10), True, RGBColor(68, 68, 68)
    
    po = doc.add_paragraph(); po.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ro = po.add_run(limpar_texto_ia(consolidacao_dict.get("orgaos_emissores") or "").replace("<br/>", "\n"))
    ro.font.name, ro.font.size, ro.bold = 'Times New Roman', Pt(11), True

    ptit = doc.add_paragraph(); ptit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = ptit.add_run(limpar_texto_ia(consolidacao_dict.get("titulo_portaria") or ""))
    rt.font.name, rt.font.size, rt.bold = 'Times New Roman', Pt(11), True

    def _render_docx_p(p_obj, texto_html, bold_all=False):
        if not texto_html: return
        for p_html in texto_html.split("<br/>"):
            if not p_html.strip(): continue
            if bold_all:
                run = p_obj.add_run(re.sub(r'<[^>]+>', '', p_html).replace("&nbsp;", "\xa0"))
                run.font.name, run.font.size, run.bold = 'Times New Roman', Pt(10), True
            else: 
                aplicar_html_no_docx(p_obj, p_html)
            p_obj.add_run("\n")

    p_ementa = doc.add_paragraph()
    p_ementa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ementa.paragraph_format.left_indent = Inches(3)
    _render_docx_p(p_ementa, consolidacao_dict.get("ementa", ""))

    p_preambulo = doc.add_paragraph()
    p_preambulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_preambulo.paragraph_format.left_indent = Inches(0)
    _render_docx_p(p_preambulo, consolidacao_dict.get("preambulo", ""))

    for item in consolidacao_dict.get("dispositivos", []):
        t = (item.get("tipo") or "").lower()
        t_prin = injetar_nota_remissiva(item.get(f"texto_principal_{tipo_versao}"), item.get("nota_remissiva") if not item.get("is_tabela") else "")
        
        if "capitulo" in t or "anexo" in t: 
            if "anexo" in t: doc.add_page_break()
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _render_docx_p(p, t_prin, bold_all=True)
            if not item.get("is_tabela"):
                continue
        else:
            if t_prin:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent = Inches(0.4)
                _render_docx_p(p, t_prin)
        
        if item.get("is_tabela"):
            linhas = item.get(f"tabela_{tipo_versao}") or []
            if linhas:
                tb = doc.add_table(rows=len(linhas), cols=len(linhas[0])); tb.style = 'Table Grid'
                for r_idx, linha in enumerate(linhas):
                    for c_idx, celula in enumerate(linha):
                        _render_docx_p(tb.cell(r_idx, c_idx).paragraphs[0], celula.replace('\n', '<br/>'))
            t_pos = injetar_nota_remissiva(item.get(f"texto_pos_tabela_{tipo_versao}"), item.get("nota_remissiva"))
            if t_pos:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent = Inches(0.4)
                _render_docx_p(p, t_pos)

    pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER; pa.paragraph_format.space_before = Pt(36)
    ra = pa.add_run(f"{limpar_texto_ia(consolidacao_dict.get('assinatura_nome') or '')}\n{limpar_texto_ia(consolidacao_dict.get('assinatura_cargo') or '')}")
    ra.font.name, ra.font.size, ra.bold = 'Times New Roman', Pt(11), True
    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer.getvalue()

def salvar_no_supabase(cons, cons_original):
    if not supabase: return False
    try:
        if cons_original:
            def _registrar(campo, original, editado):
                if original != editado and (original or editado):
                    try: supabase.table("memoria_de_correcoes").insert({"texto_ia": json.dumps(original) if not isinstance(original, str) else original, "texto_corrigido": json.dumps(editado) if not isinstance(editado, str) else editado}).execute()
                    except: pass

            _registrar("ementa", cons_original.get('ementa'), cons.get('ementa'))
            _registrar("preambulo", cons_original.get('preambulo'), cons.get('preambulo'))
            for j, disp_editado in enumerate(cons.get("dispositivos", [])):
                if j >= len(cons_original.get("dispositivos", [])): break
                disp_original = cons_original["dispositivos"][j]
                for campo in ["texto_principal_alterada", "texto_principal_consolidada", "tabela_alterada", "tabela_consolidada", "texto_pos_tabela_alterada", "texto_pos_tabela_consolidada", "nota_remissiva"]:
                    _registrar(campo, disp_original.get(campo), disp_editado.get(campo))
        
        base = cons['norma_base']
        alteradoras = cons.get('normas_alteradoras', [])
        data_base_iso = converter_para_iso(base.get('data_assinatura'))
        # upsert atômico por 'nome_padronizado' (requer a UNIQUE constraint de
        # fix_concorrencia.sql) — elimina a corrida entre dois usuários
        # salvando o mesmo ato original ao mesmo tempo, que antes podia gerar
        # linhas duplicadas com o padrão "verifica se existe, depois insere".
        res_upsert = supabase.table("portarias_base").upsert({
            "tipo_documento": base['tipo_documento'], "numero_documento": base['numero_documento'],
            "orgao_emissor": base['orgao_emissor'], "data_assinatura": data_base_iso,
            "nome_padronizado": base['nome_padronizado'], "titulo_original": cons.get("titulo_portaria"),
            "orgaos_emissores": cons.get("orgaos_emissores"), "assinatura_nome": cons.get("assinatura_nome"),
            "assinatura_cargo": cons.get("assinatura_cargo"), "documento_consolidado_json": cons,
        }, on_conflict="nome_padronizado").execute()
        if res_upsert.data:
            base_id = res_upsert.data[0]['id']
        else:
            base_id = supabase.table("portarias_base").select("id").eq("nome_padronizado", base['nome_padronizado']).execute().data[0]['id']

        for alt in alteradoras:
            res_alt = supabase.table("portarias_alteradoras").select("id").eq("portaria_base_id", base_id).eq("nome_padronizado", alt['nome_padronizado']).execute()
            if not res_alt.data:
                data_alt_iso = converter_para_iso(alt.get('data_assinatura'))
                supabase.table("portarias_alteradoras").insert({"portaria_base_id": base_id, "tipo_documento": alt['tipo_documento'], "numero_documento": alt['numero_documento'], "orgao_emissor": alt['orgao_emissor'], "data_assinatura": data_alt_iso, "nome_padronizado": alt['nome_padronizado'], "arquivo_nome_original": "Múltiplos Documentos"}).execute()
        return True
    except: return False

# ----------------- FRONT-END COM EDITOR VISUAL -----------------
if "dados_processados" not in st.session_state: st.session_state.dados_processados = None
if "dados_originais_ia" not in st.session_state: st.session_state.dados_originais_ia = None
st.markdown("<br>", unsafe_allow_html=True)

if st.button("🚀 Iniciar Análise Autopilot", type="primary", use_container_width=True):
    if not api_key: st.error("⚠️ Insira sua chave da API nas configurações.")
    elif not arquivos_enviados: st.warning("⚠️ Envie os arquivos normativos primeiro.")
    else:
        with st.spinner("⚡ Executando OCR Estrutural e Consulta ao Histórico de Aprendizado..."):
            try:
                st.session_state.dados_processados = analisar_lote_arquivos(arquivos_enviados, api_key.strip(), provedor_escolhido)
                st.session_state.dados_originais_ia = copy.deepcopy(st.session_state.dados_processados)
                st.success("✨ Processamento concluído!")
            except Exception as e:
                st.error(f"❌ Ocorreu um erro: {str(e)}")

if st.session_state.dados_processados:
    st.markdown("---")
    dados = st.session_state.dados_processados
    dados_originais = st.session_state.dados_originais_ia
    
    for i, cons in enumerate(dados.get("consolidacoes_geradas", [])):
        nome_exibicao_base = cons['norma_base']['nome_padronizado']
        nomes_alteradoras = [alt['nome_padronizado'] for alt in cons.get('normas_alteradoras', [])]
        nome_exibicao_alt = " e ".join(nomes_alteradoras) if nomes_alteradoras else "Desconhecido"
        
        with st.expander(f"📁 **{nome_exibicao_base}** alterada por **{nome_exibicao_alt}**", expanded=True):
            st.markdown("### 📝 Editor Visual de Documento")

            cons['titulo_portaria'] = st.text_input("Título do Ato Normativo", cons.get('titulo_portaria', ''), key=f"titulo_{i}")

            st.markdown("**Ementa**")
            val_ementa = ia_para_editor(cons.get('ementa', ''))
            ementa_editada = st_quill(value=val_ementa, html=True, key=f"q_ementa_{i}")
            if ementa_editada is not None: cons['ementa'] = editor_para_pdf(ementa_editada)

            st.markdown("**Preâmbulo e Considerandos**")
            val_preambulo = ia_para_editor(cons.get('preambulo', ''))
            preambulo_editado = st_quill(value=val_preambulo, html=True, key=f"q_preambulo_{i}")
            if preambulo_editado is not None: cons['preambulo'] = editor_para_pdf(preambulo_editado)

            st.markdown("#### Dispositivos (Artigos, Parágrafos, Incisos, Anexos)")
            for j, disp in enumerate(cons.get("dispositivos", [])):
                st.markdown(f"**{disp.get('tipo', 'Dispositivo').upper()} {j+1}**")
                c_alt, c_cons = st.columns(2)

                with c_alt:
                    st.markdown("*Versão Alterada*")
                    val_alt = ia_para_editor(disp.get('texto_principal_alterada', ''))
                    alt_editada = st_quill(value=val_alt, html=True, key=f"q_alt_{i}_{j}")
                    if alt_editada is not None: disp['texto_principal_alterada'] = editor_para_pdf(alt_editada)

                with c_cons:
                    st.markdown("*Versão Consolidada*")
                    val_cons = ia_para_editor(disp.get('texto_principal_consolidada', ''))
                    cons_editada = st_quill(value=val_cons, html=True, key=f"q_cons_{i}_{j}")
                    if cons_editada is not None: disp['texto_principal_consolidada'] = editor_para_pdf(cons_editada)

                st.markdown("*Nota Remissiva (Injetada automaticamente no final)*")
                nota_editada = st.text_input("Nota", value=disp.get('nota_remissiva', ''), key=f"nota_{i}_{j}", label_visibility="collapsed")
                disp['nota_remissiva'] = nota_editada
                st.markdown("---")

                if disp.get('is_tabela'):
                    st.markdown("*Tabela / Anexo — revise linha a linha*")
                    t_alt, t_cons = st.columns(2)
                    with t_alt:
                        tab_alt_edit = st.data_editor(disp.get('tabela_alterada') or [[""]], key=f"tab_alt_{i}_{j}", num_rows="dynamic", use_container_width=True)
                        disp['tabela_alterada'] = tab_alt_edit if isinstance(tab_alt_edit, list) else disp.get('tabela_alterada')
                        pos_alt = st.text_area("Texto após a tabela (Alterada)", value=disp.get('texto_pos_tabela_alterada') or "", key=f"pos_alt_{i}_{j}")
                        disp['texto_pos_tabela_alterada'] = pos_alt
                    with t_cons:
                        tab_cons_edit = st.data_editor(disp.get('tabela_consolidada') or [[""]], key=f"tab_cons_{i}_{j}", num_rows="dynamic", use_container_width=True)
                        disp['tabela_consolidada'] = tab_cons_edit if isinstance(tab_cons_edit, list) else disp.get('tabela_consolidada')
                        pos_cons = st.text_area("Texto após a tabela (Consolidada)", value=disp.get('texto_pos_tabela_consolidada') or "", key=f"pos_cons_{i}_{j}")
                        disp['texto_pos_tabela_consolidada'] = pos_cons

            st.markdown("### 📥 Opções de Exportação")
            if st.button(f"💾 Salvar Cascata Inteira no Banco de Dados", key=f"btn_sup_{i}"):
                cons_original = dados_originais.get("consolidacoes_geradas", [])[i] if dados_originais else None
                if salvar_no_supabase(cons, cons_original): st.success(f"Banco atualizado!")
            
            c_html, c_pdf, c_docx = st.columns(3)
            nome_arquivo_base = nome_exibicao_base.replace(' ', '_').replace('/', '-')

            try:
                html_alt = gerar_html_dinamico(cons, "alterada")
                html_cons = gerar_html_dinamico(cons, "consolidada")
                c_html.download_button("🌐 Baixar HTML (Alterada)", data=html_alt, file_name=f"{nome_arquivo_base}_Alt.html", mime="text/html", key=f"ha_{i}")
                c_html.download_button("🌐 Baixar HTML (Consolidada)", data=html_cons, file_name=f"{nome_arquivo_base}_Cons.html", mime="text/html", key=f"hc_{i}")
            except Exception as e:
                c_html.error(f"Falha ao gerar HTML: {e}")

            try:
                pdf_alt = gerar_pdf_dinamico(cons, "alterada")
                pdf_cons = gerar_pdf_dinamico(cons, "consolidada")
                c_pdf.download_button("📄 Baixar PDF (Alterada)", data=pdf_alt, file_name=f"{nome_arquivo_base}_Alt.pdf", mime="application/pdf", key=f"pa_{i}")
                c_pdf.download_button("📄 Baixar PDF (Consolidada)", data=pdf_cons, file_name=f"{nome_arquivo_base}_Cons.pdf", mime="application/pdf", key=f"pc_{i}")
            except Exception as e:
                c_pdf.error(f"Falha ao gerar PDF: {e}")

            try:
                docx_alt = gerar_docx_dinamico(cons, "alterada")
                docx_cons = gerar_docx_dinamico(cons, "consolidada")
                c_docx.download_button("📝 Baixar DOCX (Alterada)", data=docx_alt, file_name=f"{nome_arquivo_base}_Alt.docx", mime="application/vnd.openxmlformats", key=f"da_{i}")
                c_docx.download_button("📝 Baixar DOCX (Consolidada)", data=docx_cons, file_name=f"{nome_arquivo_base}_Cons.docx", mime="application/vnd.openxmlformats", key=f"dc_{i}")
            except Exception as e:
                c_docx.error(f"Falha ao gerar DOCX: {e}")

    if st.button("🔄 Nova Análise", type="secondary"): st.session_state.dados_processados = None; st.session_state.dados_originais_ia = None; st.rerun()
