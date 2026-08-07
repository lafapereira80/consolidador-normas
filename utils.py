import tempfile
import os
import re
import io
import json
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
from typing import List, Optional
from supabase import create_client, Client
import streamlit as st

# Importações para Word e PDF
import docx
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.graphics.shapes import Drawing, Line

def init_supabase() -> Optional[Client]:
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except:
        return None

# Pydantic Schemas
class Dispositivo(BaseModel):
    tipo: str
    texto_principal_alterada: str
    texto_principal_consolidada: str
    is_tabela: bool = False
    tabela_alterada: Optional[List[List[str]]] = None
    tabela_consolidada: Optional[List[List[str]]] = None
    texto_pos_tabela_alterada: Optional[str] = None
    texto_pos_tabela_consolidada: Optional[str] = None
    nota_remissiva: Optional[str] = ""

class Consolidacao(BaseModel):
    arquivo_original_identificado: str
    arquivo_alterador_identificado: str
    nome_portaria_base: str
    ano_portaria_base: int
    ano_portaria_alteradora: int
    nome_portaria_alteradora: str
    cabecalho_complemento: str
    orgaos_emissores: str
    titulo_portaria: str
    ementa_preambulo: str
    assinatura_nome: str
    assinatura_cargo: str
    dispositivos: List[Dispositivo]

class ArquivoAvulso(BaseModel):
    nome_arquivo: str
    nome_portaria_identificada: str
    motivo: str

class AnaliseGlobal(BaseModel):
    consolidacoes_geradas: List[Consolidacao]
    arquivos_nao_alterados: List[ArquivoAvulso]

def buscar_historico_supabase(nome_base, ano_base):
    supabase = init_supabase()
    if not supabase: return []
    try:
        res = supabase.table("portarias_base").select("id, portarias_alteradoras(nome_portaria_alteradora, ano_alteracao)").eq("nome_portaria", nome_base).eq("ano_criacao", ano_base).execute()
        if res.data and len(res.data) > 0:
            return res.data[0].get("portarias_alteradoras", [])
    except: pass
    return []

def analisar_lote_arquivos(arquivos, api_key):
    client = genai.Client(api_key=api_key)
    caminhos_temporarios, gemini_files_objs = [], []
    try:
        for arq in arquivos:
            ext = f".{arq.name.split('.')[-1]}"
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(arq.getvalue())
                caminhos_temporarios.append((tmp.name, arq.name))
                
        conteudos_prompt = ["Analise as relações entre os seguintes arquivos normativos:"]
        for caminho_tmp, nome_original in caminhos_temporarios:
            g_file = client.files.upload(file=caminho_tmp)
            gemini_files_objs.append(g_file)
            conteudos_prompt.append(f"ARQUIVO: {nome_original}")
            conteudos_prompt.append(g_file)

        conteudos_prompt.append("Atue como Especialista em Técnica Legislativa. Identifique pares, preserve formatações (b, i, strike vermelho) e notas remissivas.")

        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=conteudos_prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=AnaliseGlobal, temperature=0.0),
        )
        dados = json.loads(response.text)
        
        for cons in dados.get("consolidacoes_geradas", []):
            historico_antigo = buscar_historico_supabase(cons['nome_portaria_base'], cons['ano_portaria_base'])
            if historico_antigo:
                nomes_antigos = [h['nome_portaria_alteradora'] for h in historico_antigo]
                historico_str = " e ".join(nomes_antigos)
                cons['cabecalho_complemento'] = f"Atualizada pelas portarias {historico_str} e {cons['nome_portaria_alteradora']}"
        return dados
    finally:
        for g_file in gemini_files_objs:
            try: client.files.delete(name=g_file.name)
            except: pass
        for caminho_tmp, _ in caminhos_temporarios:
            if os.path.exists(caminho_tmp): os.remove(caminho_tmp)

def salvar_no_supabase(cons):
    supabase = init_supabase()
    if not supabase: return False, "Supabase indisponível."
    try:
        nome_base, ano_base = cons['nome_portaria_base'], cons['ano_portaria_base']
        res_busca = supabase.table("portarias_base").select("id").eq("nome_portaria", nome_base).eq("ano_criacao", ano_base).execute()
        
        if res_busca.data:
            base_id = res_busca.data[0]['id']
        else:
            res_ins = supabase.table("portarias_base").insert({
                "nome_portaria": nome_base, "ano_criacao": ano_base,
                "titulo_original": cons.get("titulo_portaria"), "orgaos_emissores": cons.get("orgaos_emissores"),
                "assinatura_nome": cons.get("assinatura_nome"), "assinatura_cargo": cons.get("assinatura_cargo")
            }).execute()
            base_id = res_ins.data[0]['id']
            
        nome_alt, ano_alt = cons['nome_portaria_alteradora'], cons['ano_portaria_alteradora']
        res_alt_check = supabase.table("portarias_alteradoras").select("id").eq("portaria_base_id", base_id).eq("nome_portaria_alteradora", nome_alt).execute()
        
        if not res_alt_check.data:
            supabase.table("portarias_alteradoras").insert({
                "portaria_base_id": base_id, "nome_portaria_alteradora": nome_alt,
                "ano_alteracao": ano_alt, "arquivo_nome_original": cons.get("arquivo_alterador_identificado")
            }).execute()
        return True, "Sucesso"
    except Exception as e:
        return False, str(e)

# Funções de texto e PDF/DOCX mantidas limpas
def extrair_paragrafos_seguros(texto_html):
    texto_html = (texto_html or "").replace("</font></strike>", "</strike></font>")
    tokens = re.split(r'(<[^>]+>)', texto_html)
    paragrafos, pilha, texto_atual = [], [], ""
    for token in tokens:
        if not token: continue
        t_lower = token.lower()
        if t_lower in ["<br>", "<br/>", "<br />"]:
            if re.sub(r'<[^>]+>', '', texto_atual).strip(): paragrafos.append(texto_atual.strip())
            texto_atual = ""
        else: texto_atual += token
    if re.sub(r'<[^>]+>', '', texto_atual).strip(): paragrafos.append(texto_atual.strip())
    return paragrafos

def injetar_nota_remissiva(texto, nota):
    if nota and nota.strip():
        n = nota.strip()
        if not n.startswith("("): n = f"({n}"
        if not n.endswith(")"): n = f"{n})"
        return f"{texto.rstrip()} &nbsp;<font color='red'>{n}</font>" if texto else f"<font color='red'>{n}</font>"
    return texto

def renderizar_paragrafos_pdf(story, texto_html, estilo):
    for p in extrair_paragrafos_seguros(texto_html):
        story.append(Paragraph(p, estilo))

def gerar_pdf_dinamico(consolidacao_dict, tipo_versao):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    story, styles = [], getSampleStyleSheet()
    estilo_dispositivo = ParagraphStyle('Dispositivo', parent=styles['Normal'], fontName='Times-Roman', fontSize=11, leading=15, alignment=4, firstLineIndent=30, spaceAfter=12)
    
    comp = consolidacao_dict.get("cabecalho_complemento", "")
    story.append(Paragraph(f"VERSÃO {tipo_versao.upper()} - {comp}", ParagraphStyle('Topo', fontName='Times-Bold', fontSize=10, alignment=1, spaceAfter=20)))
    
    preamb = (consolidacao_dict.get("ementa_preambulo") or "").replace('\n', '<br/>')
    renderizar_paragrafos_pdf(story, preamb, estilo_dispositivo)

    for item in consolidacao_dict.get("dispositivos", []):
        tipo = (item.get("tipo") or "").lower()
        nota = item.get("nota_remissiva") or ""
        texto = injetar_nota_remissiva((item.get(f"texto_principal_{tipo_versao}") or "").replace('\n', '<br/>'), nota)
        renderizar_paragrafos_pdf(story, texto, estilo_dispositivo)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_docx_dinamico(consolidacao_dict, tipo_versao):
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run(f"VERSÃO {tipo_versao.upper()} - {consolidacao_dict.get('cabecalho_complemento', '')}")
    
    for item in consolidacao_dict.get("dispositivos", []):
        nota = item.get("nota_remissiva") or ""
        texto = injetar_nota_remissiva(item.get(f"texto_principal_{tipo_versao}") or "", nota)
        doc.add_paragraph(re.sub(r'<[^>]+>', '', texto))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
